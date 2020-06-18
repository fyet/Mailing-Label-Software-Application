#########################################################################################################
# Author: Timothy Fye
# Title: mailing_label_script.py
#
# Purpose: This program merges data structures, algorithms, read/write operations required to convert an 
#   excel spread sheet containing addresses to formatted mailing lables. It builds a GUI interface so 
#   non-programmers can use this. 
#
# Description: This program implements a handful of classes & object oriented programming to house data 
#   members and functions that are pertienent to each classes operations 
#
#########################################################################################################
from tkinter import *
import tkinter as tk
from tkinter import ttk
from tkinter.filedialog import askopenfilename
from tkinter.filedialog import asksaveasfile
from tkinter import filedialog
from tkinter import messagebox
from docx import Document        # Note: Run 'pip install python-docx' in the directory in question
from docx.shared import Inches   
from docx.shared import Pt
from datetime import date
import sys
import os
import xlrd                      # Note: Run 'pip install xlrd' in the dir of sourse code if necessary
from xlrd import open_workbook   # xlrd, xlutils and xlwt modules need to be installed in order to edit worksheets
from xlutils.copy import copy    # Note: Run 'pip install xlutils' in the dir of sourse code if necessary
import xlwt                      # Note: Run 'pip install xlwt' in the dir of sourse code if necessary
import base64
#import StringIO
import io

# --- classes ---
class MyWindow:

    # Class constructor
    def __init__(self, parent):

        # Class Data Members / Class Variables

        # Local variables that are not layout related. These will be used by class methods below to store application data & process results
        self.filename = None
        self.outputFilePath = None
        self.sampleFilePath = None        
        self.data = None
        #
        self.axis_logo_base64_string = ("R0lGODlhaQBGAPcAAAAAAwAACwAAFAMDGwgIHwwMIhERJxQUKhgYLh0dMiEhNSUlOikpPS0tQTExRDc3Sjs7Tj4+UHkHZXoJZn0Pan4Ra0REVUdHWEpKW1RTY1tbamRkc2lpd2"
        "5ufHFxf4EWboQccochdYgjd4YrdIgvdokmeIwqe4kwd4s0eo46fpA8f3V1gXp6hn19iZA0gJE9gZNDg5dIh5lEi5dJiJlMipxSjp5MkJ9WkJ9YkaBQkqJclKRgl6ZjmahgnKprnq1toa9yo7F0pbJ4p7R2qrV7qoKCjoaGk"
        "YuLlY+PmZGRm5ubpJ+fqLeBrbmErruFsb2Ls6SkrKurs7W1u7i4v8GPuMCRt8OVusWavby8wsqewcukw86qx8+ryNGtytSzzde50dm91MHBx8XFy8fMzcnJzs3N0tzB1t7G2dHR1dbW2tra3eLL3d7e4ebT4ujX5enZ5uzd"
        "6eHh5Obm6O7i7Ovr7fHm7+/v8fTs8vjp9PDw8vfx9vb29/Xy9fjz9/jy9vf3+Pf3+Pn1+Pr2+fn1+fv4+vv5+/r6+vr6+/v7+/n5+vz5+/z6+/v7/Pz6/P37/P37/fz8/Pz8/f38/f39/f78/f79/f79/v7+/v/+/v7+///"
        "+/////////v///wAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
        "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
        "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAACH5BAAAAAAAIf8LTUdLOEJJTTAwMDD/OEJJTQQlAAAAAAAQAAAAAAAAAAAAAAAAAAAAADhCSU0D6gAAAAAYIDw/eG1sIHZl"
        "cnNpb249IjEuMCIgZW5jb2Rpbmc9IlVURi04Ij8+CjwhRE9DVFlQRSBwbGlzdCBQVUJMSUMgIi0vL0FwcGxlLy9EVEQgUExJU1QgMS4wLy9FTiIgImh0dHA6Ly93d3cuYXBwbGUuY29tL0RURHMvUHJvcGVydHlMaXN0LTE"
        "uMC5kdGQiPgo8cGxpc3QgdmVyc2lvbj0iMS4wIj4KPGRpY3Q+Cgk8a2V5PmNvbS5hcHBsZS5wcmludC5QYWdlRm9ybWF0LlBNSG9yaXpvbnRh/2xSZXM8L2tleT4KCTxkaWN0PgoJCTxrZXk+Y29tLmFwcGxlLnByaW50Ln"
        "RpY2tldC5jcmVhdG9yPC9rZXk+CgkJPHN0cmluZz5jb20uYXBwbGUuam9idGlja2V0PC9zdHJpbmc+CgkJPGtleT5jb20uYXBwbGUucHJpbnQudGlja2V0Lml0ZW1BcnJheTwva2V5PgoJCTxhcnJheT4KCQkJPGRpY3Q+C"
        "gkJCQk8a2V5PmNvbS5hcHBsZS5wcmludC5QYWdlRm9ybWF0LlBNSG9yaXpvbnRhbFJlczwva2V5PgoJCQkJPHJlYWw+NzI8L3JlYWw+CgkJCQk8a2V5Pv9jb20uYXBwbGUucHJpbnQudGlja2V0LnN0YXRlRmxhZzwva2V5"
        "PgoJCQkJPGludGVnZXI+MDwvaW50ZWdlcj4KCQkJPC9kaWN0PgoJCTwvYXJyYXk+Cgk8L2RpY3Q+Cgk8a2V5PmNvbS5hcHBsZS5wcmludC5QYWdlRm9ybWF0LlBNT3JpZW50YXRpb248L2tleT4KCTxkaWN0PgoJCTxrZXk"
        "+Y29tLmFwcGxlLnByaW50LnRpY2tldC5jcmVhdG9yPC9rZXk+CgkJPHN0cmluZz5jb20uYXBwbGUuam9idGlja2V0PC9zdHJpbmc+CgkJPGtleT5jb20uYXBwbGX/LnByaW50LnRpY2tldC5pdGVtQXJyYXk8L2tleT4KCQ"
        "k8YXJyYXk+CgkJCTxkaWN0PgoJCQkJPGtleT5jb20uYXBwbGUucHJpbnQuUGFnZUZvcm1hdC5QTU9yaWVudGF0aW9uPC9rZXk+CgkJCQk8aW50ZWdlcj4xPC9pbnRlZ2VyPgoJCQkJPGtleT5jb20uYXBwbGUucHJpbnQud"
        "Glja2V0LnN0YXRlRmxhZzwva2V5PgoJCQkJPGludGVnZXI+MDwvaW50ZWdlcj4KCQkJPC9kaWN0PgoJCTwvYXJyYXk+Cgk8L2RpY3Q+Cgk8a2V5PmNvbS5hcHBsZS5wcmlu/3QuUGFnZUZvcm1hdC5QTVNjYWxpbmc8L2tl"
        "eT4KCTxkaWN0PgoJCTxrZXk+Y29tLmFwcGxlLnByaW50LnRpY2tldC5jcmVhdG9yPC9rZXk+CgkJPHN0cmluZz5jb20uYXBwbGUuam9idGlja2V0PC9zdHJpbmc+CgkJPGtleT5jb20uYXBwbGUucHJpbnQudGlja2V0Lml"
        "0ZW1BcnJheTwva2V5PgoJCTxhcnJheT4KCQkJPGRpY3Q+CgkJCQk8a2V5PmNvbS5hcHBsZS5wcmludC5QYWdlRm9ybWF0LlBNU2NhbGluZzwva2V5PgoJCQkJPHJlYWw+MTwvcmVhbP8+CgkJCQk8a2V5PmNvbS5hcHBsZS"
        "5wcmludC50aWNrZXQuc3RhdGVGbGFnPC9rZXk+CgkJCQk8aW50ZWdlcj4wPC9pbnRlZ2VyPgoJCQk8L2RpY3Q+CgkJPC9hcnJheT4KCTwvZGljdD4KCTxrZXk+Y29tLmFwcGxlLnByaW50LlBhZ2VGb3JtYXQuUE1WZXJ0a"
        "WNhbFJlczwva2V5PgoJPGRpY3Q+CgkJPGtleT5jb20uYXBwbGUucHJpbnQudGlja2V0LmNyZWF0b3I8L2tleT4KCQk8c3RyaW5nPmNvbS5hcHBsZS5qb2J0aWNrZXQ8L3N0cmluZz4KCQk8a2X/eT5jb20uYXBwbGUucHJp"
        "bnQudGlja2V0Lml0ZW1BcnJheTwva2V5PgoJCTxhcnJheT4KCQkJPGRpY3Q+CgkJCQk8a2V5PmNvbS5hcHBsZS5wcmludC5QYWdlRm9ybWF0LlBNVmVydGljYWxSZXM8L2tleT4KCQkJCTxyZWFsPjcyPC9yZWFsPgoJCQk"
        "JPGtleT5jb20uYXBwbGUucHJpbnQudGlja2V0LnN0YXRlRmxhZzwva2V5PgoJCQkJPGludGVnZXI+MDwvaW50ZWdlcj4KCQkJPC9kaWN0PgoJCTwvYXJyYXk+Cgk8L2RpY3Q+Cgk8a2V5PmNvbS5hcHBs/2UucHJpbnQuUG"
        "FnZUZvcm1hdC5QTVZlcnRpY2FsU2NhbGluZzwva2V5PgoJPGRpY3Q+CgkJPGtleT5jb20uYXBwbGUucHJpbnQudGlja2V0LmNyZWF0b3I8L2tleT4KCQk8c3RyaW5nPmNvbS5hcHBsZS5qb2J0aWNrZXQ8L3N0cmluZz4KC"
        "Qk8a2V5PmNvbS5hcHBsZS5wcmludC50aWNrZXQuaXRlbUFycmF5PC9rZXk+CgkJPGFycmF5PgoJCQk8ZGljdD4KCQkJCTxrZXk+Y29tLmFwcGxlLnByaW50LlBhZ2VGb3JtYXQuUE1WZXJ0aWNhbFNjYWxpbmc8L/9rZXk+"
        "CgkJCQk8cmVhbD4xPC9yZWFsPgoJCQkJPGtleT5jb20uYXBwbGUucHJpbnQudGlja2V0LnN0YXRlRmxhZzwva2V5PgoJCQkJPGludGVnZXI+MDwvaW50ZWdlcj4KCQkJPC9kaWN0PgoJCTwvYXJyYXk+Cgk8L2RpY3Q+Cgk"
        "8a2V5PmNvbS5hcHBsZS5wcmludC5zdWJUaWNrZXQucGFwZXJfaW5mb190aWNrZXQ8L2tleT4KCTxkaWN0PgoJCTxrZXk+UE1QUERQYXBlckNvZGVOYW1lPC9rZXk+CgkJPGRpY3Q+CgkJCTxrZXk+Y29tLmFwcGxlLnByaW"
        "7/dC50aWNrZXQuY3JlYXRvcjwva2V5PgoJCQk8c3RyaW5nPmNvbS5hcHBsZS5qb2J0aWNrZXQ8L3N0cmluZz4KCQkJPGtleT5jb20uYXBwbGUucHJpbnQudGlja2V0Lml0ZW1BcnJheTwva2V5PgoJCQk8YXJyYXk+CgkJC"
        "Qk8ZGljdD4KCQkJCQk8a2V5PlBNUFBEUGFwZXJDb2RlTmFtZTwva2V5PgoJCQkJCTxzdHJpbmc+TGV0dGVyPC9zdHJpbmc+CgkJCQkJPGtleT5jb20uYXBwbGUucHJpbnQudGlja2V0LnN0YXRlRmxhZzwva2V5PgoJCQkJ"
        "CTxpbnRl/2dlcj4wPC9pbnRlZ2VyPgoJCQkJPC9kaWN0PgoJCQk8L2FycmF5PgoJCTwvZGljdD4KCQk8a2V5PlBNVGlvZ2FQYXBlck5hbWU8L2tleT4KCQk8ZGljdD4KCQkJPGtleT5jb20uYXBwbGUucHJpbnQudGlja2V"
        "0LmNyZWF0b3I8L2tleT4KCQkJPHN0cmluZz5jb20uYXBwbGUuam9idGlja2V0PC9zdHJpbmc+CgkJCTxrZXk+Y29tLmFwcGxlLnByaW50LnRpY2tldC5pdGVtQXJyYXk8L2tleT4KCQkJPGFycmF5PgoJCQkJPGRpY3Q+Cg"
        "kJCQkJPGtleT5QTf9UaW9nYVBhcGVyTmFtZTwva2V5PgoJCQkJCTxzdHJpbmc+bmEtbGV0dGVyPC9zdHJpbmc+CgkJCQkJPGtleT5jb20uYXBwbGUucHJpbnQudGlja2V0LnN0YXRlRmxhZzwva2V5PgoJCQkJCTxpbnRlZ"
        "2VyPjA8L2ludGVnZXI+CgkJCQk8L2RpY3Q+CgkJCTwvYXJyYXk+CgkJPC9kaWN0PgoJCTxrZXk+Y29tLmFwcGxlLnByaW50LlBhZ2VGb3JtYXQuUE1BZGp1c3RlZFBhZ2VSZWN0PC9rZXk+CgkJPGRpY3Q+CgkJCTxrZXk+"
        "Y29tLmFwcGxlLnByaW50LnT/aWNrZXQuY3JlYXRvcjwva2V5PgoJCQk8c3RyaW5nPmNvbS5hcHBsZS5qb2J0aWNrZXQ8L3N0cmluZz4KCQkJPGtleT5jb20uYXBwbGUucHJpbnQudGlja2V0Lml0ZW1BcnJheTwva2V5Pgo"
        "JCQk8YXJyYXk+CgkJCQk8ZGljdD4KCQkJCQk8a2V5PmNvbS5hcHBsZS5wcmludC5QYWdlRm9ybWF0LlBNQWRqdXN0ZWRQYWdlUmVjdDwva2V5PgoJCQkJCTxhcnJheT4KCQkJCQkJPGludGVnZXI+MDwvaW50ZWdlcj4KCQ"
        "kJCQkJPGludGVnZXI+MDwvaW50ZWdl/3I+CgkJCQkJCTxyZWFsPjczNDwvcmVhbD4KCQkJCQkJPHJlYWw+NTc2PC9yZWFsPgoJCQkJCTwvYXJyYXk+CgkJCQkJPGtleT5jb20uYXBwbGUucHJpbnQudGlja2V0LnN0YXRlR"
        "mxhZzwva2V5PgoJCQkJCTxpbnRlZ2VyPjA8L2ludGVnZXI+CgkJCQk8L2RpY3Q+CgkJCTwvYXJyYXk+CgkJPC9kaWN0PgoJCTxrZXk+Y29tLmFwcGxlLnByaW50LlBhZ2VGb3JtYXQuUE1BZGp1c3RlZFBhcGVyUmVjdDwv"
        "a2V5PgoJCTxkaWN0PgoJCQk8a2V5PmNvbS5hcP9wbGUucHJpbnQudGlja2V0LmNyZWF0b3I8L2tleT4KCQkJPHN0cmluZz5jb20uYXBwbGUuam9idGlja2V0PC9zdHJpbmc+CgkJCTxrZXk+Y29tLmFwcGxlLnByaW50LnR"
        "pY2tldC5pdGVtQXJyYXk8L2tleT4KCQkJPGFycmF5PgoJCQkJPGRpY3Q+CgkJCQkJPGtleT5jb20uYXBwbGUucHJpbnQuUGFnZUZvcm1hdC5QTUFkanVzdGVkUGFwZXJSZWN0PC9rZXk+CgkJCQkJPGFycmF5PgoJCQkJCQ"
        "k8cmVhbD4tMTg8L3JlYWw+CgkJCQkJCTxyZWFsPi0xODz/L3JlYWw+CgkJCQkJCTxyZWFsPjc3NDwvcmVhbD4KCQkJCQkJPHJlYWw+NTk0PC9yZWFsPgoJCQkJCTwvYXJyYXk+CgkJCQkJPGtleT5jb20uYXBwbGUucHJpb"
        "nQudGlja2V0LnN0YXRlRmxhZzwva2V5PgoJCQkJCTxpbnRlZ2VyPjA8L2ludGVnZXI+CgkJCQk8L2RpY3Q+CgkJCTwvYXJyYXk+CgkJPC9kaWN0PgoJCTxrZXk+Y29tLmFwcGxlLnByaW50LlBhcGVySW5mby5QTVBhcGVy"
        "TmFtZTwva2V5PgoJCTxkaWN0PgoJCQk8a2V5PmNvbS5hcHBsZS5w/3JpbnQudGlja2V0LmNyZWF0b3I8L2tleT4KCQkJPHN0cmluZz5jb20uYXBwbGUuam9idGlja2V0PC9zdHJpbmc+CgkJCTxrZXk+Y29tLmFwcGxlLnB"
        "yaW50LnRpY2tldC5pdGVtQXJyYXk8L2tleT4KCQkJPGFycmF5PgoJCQkJPGRpY3Q+CgkJCQkJPGtleT5jb20uYXBwbGUucHJpbnQuUGFwZXJJbmZvLlBNUGFwZXJOYW1lPC9rZXk+CgkJCQkJPHN0cmluZz5uYS1sZXR0ZX"
        "I8L3N0cmluZz4KCQkJCQk8a2V5PmNvbS5hcHBsZS5wcmludC50aWNrZXQuc/90YXRlRmxhZzwva2V5PgoJCQkJCTxpbnRlZ2VyPjA8L2ludGVnZXI+CgkJCQk8L2RpY3Q+CgkJCTwvYXJyYXk+CgkJPC9kaWN0PgoJCTxrZ"
        "Xk+Y29tLmFwcGxlLnByaW50LlBhcGVySW5mby5QTVVuYWRqdXN0ZWRQYWdlUmVjdDwva2V5PgoJCTxkaWN0PgoJCQk8a2V5PmNvbS5hcHBsZS5wcmludC50aWNrZXQuY3JlYXRvcjwva2V5PgoJCQk8c3RyaW5nPmNvbS5h"
        "cHBsZS5qb2J0aWNrZXQ8L3N0cmluZz4KCQkJPGtleT5jb20uYXBwbGUucHJpbnQudGn/Y2tldC5pdGVtQXJyYXk8L2tleT4KCQkJPGFycmF5PgoJCQkJPGRpY3Q+CgkJCQkJPGtleT5jb20uYXBwbGUucHJpbnQuUGFwZXJ"
        "JbmZvLlBNVW5hZGp1c3RlZFBhZ2VSZWN0PC9rZXk+CgkJCQkJPGFycmF5PgoJCQkJCQk8aW50ZWdlcj4wPC9pbnRlZ2VyPgoJCQkJCQk8aW50ZWdlcj4wPC9pbnRlZ2VyPgoJCQkJCQk8cmVhbD43MzQ8L3JlYWw+CgkJCQ"
        "kJCTxyZWFsPjU3NjwvcmVhbD4KCQkJCQk8L2FycmF5PgoJCQkJCTxrZXk+Y29tLmFwcGxlLnBy/2ludC50aWNrZXQuc3RhdGVGbGFnPC9rZXk+CgkJCQkJPGludGVnZXI+MDwvaW50ZWdlcj4KCQkJCTwvZGljdD4KCQkJP"
        "C9hcnJheT4KCQk8L2RpY3Q+CgkJPGtleT5jb20uYXBwbGUucHJpbnQuUGFwZXJJbmZvLlBNVW5hZGp1c3RlZFBhcGVyUmVjdDwva2V5PgoJCTxkaWN0PgoJCQk8a2V5PmNvbS5hcHBsZS5wcmludC50aWNrZXQuY3JlYXRv"
        "cjwva2V5PgoJCQk8c3RyaW5nPmNvbS5hcHBsZS5qb2J0aWNrZXQ8L3N0cmluZz4KCQkJPGtleT5jb20uYf9wcGxlLnByaW50LnRpY2tldC5pdGVtQXJyYXk8L2tleT4KCQkJPGFycmF5PgoJCQkJPGRpY3Q+CgkJCQkJPGt"
        "leT5jb20uYXBwbGUucHJpbnQuUGFwZXJJbmZvLlBNVW5hZGp1c3RlZFBhcGVyUmVjdDwva2V5PgoJCQkJCTxhcnJheT4KCQkJCQkJPHJlYWw+LTE4PC9yZWFsPgoJCQkJCQk8cmVhbD4tMTg8L3JlYWw+CgkJCQkJCTxyZW"
        "FsPjc3NDwvcmVhbD4KCQkJCQkJPHJlYWw+NTk0PC9yZWFsPgoJCQkJCTwvYXJyYXk+CgkJCQkJPGtleT5jb20uYXD/cGxlLnByaW50LnRpY2tldC5zdGF0ZUZsYWc8L2tleT4KCQkJCQk8aW50ZWdlcj4wPC9pbnRlZ2VyP"
        "goJCQkJPC9kaWN0PgoJCQk8L2FycmF5PgoJCTwvZGljdD4KCQk8a2V5PmNvbS5hcHBsZS5wcmludC5QYXBlckluZm8ucHBkLlBNUGFwZXJOYW1lPC9rZXk+CgkJPGRpY3Q+CgkJCTxrZXk+Y29tLmFwcGxlLnByaW50LnRp"
        "Y2tldC5jcmVhdG9yPC9rZXk+CgkJCTxzdHJpbmc+Y29tLmFwcGxlLmpvYnRpY2tldDwvc3RyaW5nPgoJCQk8a2V5PmNvbS5h/3BwbGUucHJpbnQudGlja2V0Lml0ZW1BcnJheTwva2V5PgoJCQk8YXJyYXk+CgkJCQk8ZGl"
        "jdD4KCQkJCQk8a2V5PmNvbS5hcHBsZS5wcmludC5QYXBlckluZm8ucHBkLlBNUGFwZXJOYW1lPC9rZXk+CgkJCQkJPHN0cmluZz5VUyBMZXR0ZXI8L3N0cmluZz4KCQkJCQk8a2V5PmNvbS5hcHBsZS5wcmludC50aWNrZX"
        "Quc3RhdGVGbGFnPC9rZXk+CgkJCQkJPGludGVnZXI+MDwvaW50ZWdlcj4KCQkJCTwvZGljdD4KCQkJPC9hcnJheT4KCQk8L2RpY3Q+Cv8JCTxrZXk+Y29tLmFwcGxlLnByaW50LnRpY2tldC5BUElWZXJzaW9uPC9rZXk+C"
        "gkJPHN0cmluZz4wMC4yMDwvc3RyaW5nPgoJCTxrZXk+Y29tLmFwcGxlLnByaW50LnRpY2tldC50eXBlPC9rZXk+CgkJPHN0cmluZz5jb20uYXBwbGUucHJpbnQuUGFwZXJJbmZvVGlja2V0PC9zdHJpbmc+Cgk8L2RpY3Q+"
        "Cgk8a2V5PmNvbS5hcHBsZS5wcmludC50aWNrZXQuQVBJVmVyc2lvbjwva2V5PgoJPHN0cmluZz4wMC4yMDwvc3RyaW5nPgoJPGtleT5jb20uYXD/cGxlLnByaW50LnRpY2tldC50eXBlPC9rZXk+Cgk8c3RyaW5nPmNvbS5"
        "hcHBsZS5wcmludC5QYWdlRm9ybWF0VGlja2V0PC9zdHJpbmc+CjwvZGljdD4KPC9wbGlzdD4KOEJJTQPtAAAAAAAQASwAAAABAAEBLAAAAAEAAThCSU0EJgAAAAAADgAAAAAAAAAAAAA/gAAAOEJJTQQNAAAAAAAEAAAAeD"
        "hCSU0EGQAAAAAABAAAAB44QklNA/MAAAAAAAkAAAAAAAAAAAEAOEJJTQQKAAAAAAABAAA4QklNJxAAAAAAAAoAAQAAAAAAAAABOEJJTQP1AAAAAABIAC9m/2YAAQBsZmYABgAAAAAAAQAvZmYAAQChmZoABgAAAAAAAQAyA"
        "AAAAQBaAAAABgAAAAAAAQA1AAAAAQAtAAAABgAAAAAAAThCSU0D+AAAAAAAcAAA/////////////////////////////wPoAAAAAP////////////////////////////8D6AAAAAD/////////////////////////////"
        "A+gAAAAA/////////////////////////////wPoAAA4QklNBAAAAAAAAAIAAThCSU0EAgAAAAAABAAAAAA4QklNBDAAAAAAAAIBAThCSU0ELQAAAAAABgABAAAAAv84QklNBAgAAAAAABAAAAABAAACQAAAAkAAAAAAOEJ"
        "JTQQeAAAAAAAEAAAAADhCSU0EGgAAAAADSQAAAAYAAAAAAAAAAAAAArQAAAQaAAAACgBVAG4AdABpAHQAbABlAGQALQAzAAAAAQAAAAAAAAAAAAAAAAAAAAAAAAABAAAAAAAAAAAAAAQaAAACtAAAAAAAAAAAAAAAAAAAAA"
        "ABAAAAAAAAAAAAAAAAAAAAAAAAABAAAAABAAAAAAAAbnVsbAAAAAIAAAAGYm91bmRzT2JqYwAAAAEAAAAAAABSY3QxAAAABAAAAABUb3AgbG9uZwAAAAAAAAAATGVmdGxvbmf/AAAAAAAAAABCdG9tbG9uZwAAArQAAAAAU"
        "mdodGxvbmcAAAQaAAAABnNsaWNlc1ZsTHMAAAABT2JqYwAAAAEAAAAAAAVzbGljZQAAABIAAAAHc2xpY2VJRGxvbmcAAAAAAAAAB2dyb3VwSURsb25nAAAAAAAAAAZvcmlnaW5lbnVtAAAADEVTbGljZU9yaWdpbgAAAA1h"
        "dXRvR2VuZXJhdGVkAAAAAFR5cGVlbnVtAAAACkVTbGljZVR5cGUAAAAASW1nIAAAAAZib3VuZHNPYmpjAAAAAQAAAAAAAFJjdDEAAAAEAAAAAFRvcCBsb25nAAAAAAAAAABMZWZ0bG9u/2cAAAAAAAAAAEJ0b21sb25nAAA"
        "CtAAAAABSZ2h0bG9uZwAABBoAAAADdXJsVEVYVAAAAAEAAAAAAABudWxsVEVYVAAAAAEAAAAAAABNc2dlVEVYVAAAAAEAAAAAAAZhbHRUYWdURVhUAAAAAQAAAAAADmNlbGxUZXh0SXNIVE1MYm9vbAEAAAAIY2VsbFRleH"
        "RURVhUAAAAAQAAAAAACWhvcnpBbGlnbmVudW0AAAAPRVNsaWNlSG9yekFsaWduAAAAB2RlZmF1bHQAAAAJdmVydEFsaWduZW51bQAAAA9FU2xpY2VWZXJ0QWxpZ24AAAAHZGVmYXVsdAAAAAtiZ/9Db2xvclR5cGVlbnVtA"
        "AAAEUVTbGljZUJHQ29sb3JUeXBlAAAAAE5vbmUAAAAJdG9wT3V0c2V0bG9uZwAAAAAAAAAKbGVmdE91dHNldGxvbmcAAAAAAAAADGJvdHRvbU91dHNldGxvbmcAAAAAAAAAC3JpZ2h0T3V0c2V0bG9uZwAAAAAAOEJJTQQo"
        "AAAAAAAMAAAAAT/wAAAAAAAAOEJJTQQUAAAAAAAEAAAAAzhCSU0EDAAAAAAQNQAAAAEAAACgAAAAaQAAAeAAAMTgAAAQGQAYAAH/2P/gABBKRklGAAECAABIAEgAAP/tAAxBZG9iZV9DTQAB/+4ADkFkb2L/ZQBkgAAAAAH"
        "/2wCEAAwICAgJCAwJCQwRCwoLERUPDAwPFRgTExUTExgRDAwMDAwMEQwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwBDQsLDQ4NEA4OEBQODg4UFA4ODg4UEQwMDAwMEREMDAwMDAwRDAwMDAwMDAwMDAwMDAwMDAwMDA"
        "wMDAwMDAwMDP/AABEIAGkAoAMBIgACEQEDEQH/3QAEAAr/xAE/AAABBQEBAQEBAQAAAAAAAAADAAECBAUGBwgJCgsBAAEFAQEBAQEBAAAAAAAAAAEAAgMEBQYHCAkKCxAAAQQBAwIEAgUHBggFAwwzAQACEQMEIRIx/wVBU"
        "WETInGBMgYUkaGxQiMkFVLBYjM0coLRQwclklPw4fFjczUWorKDJkSTVGRFwqN0NhfSVeJl8rOEw9N14/NGJ5SkhbSVxNTk9KW1xdXl9VZmdoaWprbG1ub2N0dXZ3eHl6e3x9fn9xEAAgIBAgQEAwQFBgcHBgU1AQACEQMh"
        "MRIEQVFhcSITBTKBkRShsUIjwVLR8DMkYuFygpJDUxVjczTxJQYWorKDByY1wtJEk1SjF2RFVTZ0ZeLys4TD03Xj80aUpIW0lcTU5PSltcXV5fVWZnaGlqa2xtbm9ic3R1dnd4eXp7fH/9oADAMBAAIRAxEAPwD1VJJJJSk"
        "kkv9JSkkkklKSSSSUpJJJJSkkkI5WMMkYhtYMlzDY2ncN5YDsNnp/S2bj9JJSOy3OLopobt/ets2/5rKmXbv86tRNPUniDk11j/g6jI/tW22N/wDAlbSRtN+Ac6zpebYI/a2Wyf3G4w/6rDe7/pKi/wCrnVq3F+H9YMxjiZ"
        "/WG1Xt/wC2/TpW+klxHw+xIySG1f4sXmLrvr306HmvE6xQ36QqDqbiB+dsc417v+K9b/i1Pp3176Tk3fZs1lnTcoHY9mQIYH/uOt/wX/oSzHXSLL639XOm9aq25TNt7RtqyWQLG/yZ/wAJX/wNn6NOBifmH1ivE4S0nGv60"
        "NP+b8rqJLj/fpnVeo/VTPZ0XrbvU6c/+iZeu1jZ2hzXO+jjs+hdS/8AoP8A4U9N67hCUa8QdityYzAjW4nWMh+k/wD/0PVUklhfXT6xs+rf1fyOoAg5Lv0OGw6g3PB9P+zU1r73/wAipJTupLxfpv13+ufRc/pOf9YL7buk"
        "9RabNr2NG6kuNLrW7at/q0/o8tjGfzlL6f8AuQvZmPZYxtlbg9jwHNc0yCDq1zXBEilMkklg/WnB6/bj/a+hZtlORU334o2FtrRJ/R+qx+zI/wChYkBZq6XRHEQLEb6nZ3klwv1N+ut1+QOk9bsJyHvLcbIeAwl8x9kvaAz"
        "Zdu/mfb/wH89s9Xuk/2UTE0U5McscuGX++pJeef4xf8YGT025vQPq+7d1SwtF9zBvdXujZj0Vw7flXf8AgTP+Gs/QdH9TunfWHF6eMj6x51mX1DIAJoO0V0N5FX6JrfVv/wBNb/1qn/S3ilj0C5+z6nYVn1mb1991rnNh4x"
        "yTHqtGxlgu3ep6La/+0v8ANf8AW/0S6BeWdW+tX1ho/wAZ7OjVZz2dOOZi1nHAZt2WModaySzf7970ok9CujOUbo1Yo+T6mkkuW/xjfWfL+rfQBk4Ib9rybRj1PdqK9zbLHX+mfbY5npexr/Zvf+f/ADaC16lJeT4lH+OTq"
        "GJRn0ZwNOXWy+o7qGeyxosZ7P/027Pa5F/Zf+Or/uaP+3KP/II14qfU0l5F1O//ABu9Awbeq52c0Y1BZ6kmiz6bm1N/R+n++/8ANXo31T627r/1ewurWMFVmSw+qxv0Q+tzqLdklzvTdZU51fu+gkQpL1/otHWunPxLIbaP"
        "fj2/uWAex39T8y1n59SxvqJ1e63Ht6Lmy3L6dLWNcZd6bT6Tqz+99kt/Q/8AFfZ11a4frLf2P9e8LPr9tWfsbb2BLyMO/wDzd2Jf/wAYnR1Bj9R5s2L1RljPbih/ej/3z//R9VXkPXrX/wCMD/GBR0bHcXdJ6aXNssYRGxh"
        "b9uyGvAd/PWCvEof76/5m3/CLsv8AGT//Wg/V76vubjv2dQz5oxSDDmiP1jJbDmO/QVu9j2/QvtoXmn1I+uP/ADToyNvRn5uTluaXZPqur/RtH6KprPs9/wCc6yzfu/Sf9bTgNLU+n/X36q1dc+rL8XEqa3KwR6uAxjQNWD"
        "acVgA9rLqf0TK/oer6H+jWN/ii+tJ6l0p3Q8p85fTWg0EnV+MTtZy539Ff+h/4n7MqH/j05X/zvWf+xDv/AHjXED6y3YH1u/5y4GE7AY+42uw3OLmuDx+uU+q+pntv32f4H9X9T9H/ADbEgCRSn6HSVfAzsbqGFRnYj/Ux8"
        "ljbKn+LXDcJH5rv3mqwmqeQ+uf1KZ1Vr+odOYG9/0AP0tWgbeB/0WZLf8Hb+f8Azdv5llWX0T63dZyenX9BNzaOvbXVdOy8tphz26HHy2u+jnVNa/0n2s/S2fz9dlu+vI1Prr9bz08HpXTHF3UbQBY9nuNIf9BrGj6WXbu/"
        "Q1/mfzn+i9Sr0n6jdRxeh5V1WQMX6x5lRFGQ73jH3e51Qs9zvtF7fZk5rP0lW/8AQep6X6ea/wBX6/8AA/ebV/qP1v8A1L9//wBEQ/4vf8X93S7D1zrzfU6vaXOrre4WGrcffdZZL/Uyrv3936P/AIxd+vJvqN9esz6v5Z+"
        "qv1qDqK6HelRfbzQ782jIf+fhu/wGR/gf/Cez7N6wCCJGoP/wVEbvVqrrxbrn/wCWWv8A8PYf/UYy9pXi3XP/AMstf/h7D/6jGSjup9pXnv8Ajqqe76t4dgEtZmNDvLdXdB/BehLI+snVfqzh4T8f6xX47ca9snHvh5eGmd"
        "zMYB9tmx4/wdfsekN1PM/V/wDxlfU7C6D03Dyc1zb8bEopuaKbTD662V2N3Nq2u2ub+ar/AP46v1H/AO5zv+2Lv/SSJ/4131FOv7M5/wCHyB/6PTf+Nd9RP/Kz/wAHyP8A3oS0U1c7/GD/AIueq45wuoZXrYz3Mc6uym/YS"
        "xzba9+2r6PqMb9L9H++utwKsCrDpb01lNeEW76G44aKtr/0m+n0f0f/ss3+p7F579ffqF9U+kfVPO6j07B9DLo9H07PWufG+6mp/stusr/m7H/mLc/xU2Of9SMIO1Fb7mt+Hq2O/wC/JdFPXLiv8Y/td0u4fSY+2Pl6Vn/V"
        "VrtVw/1+f9o6t0np7dXOJkf8dZTj1/8AUWp2P5x9Wbl/50eF/wDRf//ShldH619dP8YDbupYGTi9DxDDDkVWVMdRSfoA2so/SZ9zvc3+eros/wC669YSSRJUpYX11+rrfrH9Xsjp7QPtLf02G46RcwH0+7W/pWufQ7d/pVu"
        "pIKfOv8U2R1/Cou6F1bp+XjY7Jvwr76bGMbJ/T4xssDWt3Pd69Lf/AAz/AMGu/5vrZ1nq2BjDH6NhX5WdeDF1dTrK6W8eo47TXZf/AKGj/rt36P8AR3b6SIIuyL8F0CBIEjiroXiPqR9ULaHjrfWGuOY8l+PRbJewu+nlZW"
        "/3Oy7Z/O/mf+Os/Q9ukkjKRkbKcmSWSXFL/eeP/wAYP1Fq+suH9qxA2vq+O39C8wBa0a/ZrXf+ebP8G/8AkLB/xd9b+tnSzX0TrnSs92BIZi5Rx7SaNdoqtds92J+4/wD7T/8AEfzPpySF6UsUvI+sdE61b/jYr6hX0/Jfh"
        "DNxHnKbS81bWsxxY/1tvp7GbXb/AHL1xJIGlKXn3+Nj6pdV63VhZ3SqTk24gsZfS0gPLP9+17H1teW79jq3t2M/S/pV6CkgDSnyg/Xn/GtMfsA6f90sn/0ol/z6/wAa3/lA7/2Byf8A0ovV0kbHZT451frX+M36ydOt6Lld"
        "CeynLLA54xbqiNj2XN/TZFnos99bd29ei/UfomR0L6sYXTsuBlMDn3hpkB1j3XenuHtd6TX+l7VvJJEqUuBrd+3P8YO9vvx8Fxg9g3F9k/2uoWrpvrT1n9kdJstrdGVf+hxR/wAI4fzn9Whm65ZX+LzpP2bp1nUbBD82G0z"
        "z6Nchjv8Ar9jrLf5dfop8dImXf0hnx+jHPIeo4I/X5n//0/VUkkklKSSSSUpQba11j6x9KuJ+Y3L/msjKz2dP67W3Idsx8/GdsdqYtxS657fb/pca97v/AEFRAtMRdgdrddJMCCJGoKdBCkkkklKSSSSUpJJJJSkkkklKUL"
        "baqan3XOFdVbS+x7jAa1o3Oc4/utane9lbHPe4MYwFznOMAAalziV559Yuv5X1kzK+kdIaX4rnw0DT13N93rWfuYVH85/4L/oU6MTI+HUsmLEZntEfNLswsdk/XX6xtY0Pr6dSPga8efe8/uZOc5vt/wBGz/wtYvRq666q2"
        "1VtDK2ANYxogAAQ1rQs7oHQ8fomAMas77n+/JvIgvsjn+TWz6FVf5jFpozleg+UbJzZBIiMdIR0j/3z/9T1/1SSSSUpJJJJSllfWPpt2f08OxY+3YdjcrDng2V6+k7+RfXvp/64tVJEGjaYkxII6PLfUr6x05uK3pmQfTzM"
        "YFtTHaF9TdGBs/4XGZ+hvZ/Ofo/VXUrjPrb9VLzeetdGa4ZDT6l9NWjy8f8AavFj/tR/pav+1H/Hb/tEvq99fce9jMfrLm0W8NzBpU8zt/Tf9xbP3936D/iv5lPlHi9UfqOzNPGJj3Met/ND9KJexSUWPY9oexwcxwBa4GQ"
        "QeCCpKNgUkkkkpSSSBmZ2Hg1G7Mvrx6h+fY4NHwG76TklAXoE6q9R6ng9Mxjk51zaagYBOpcf3K2Nl9j/AOQxct1j/P9iY9bXM6TV6p4+1XgsrH9Sn2XW/wBv7Osvp31a699Y8kZ/VbbKaSP5+4Ra5p/MxMYgMx6/5b69n/"
        "BZCkGPrL0j8WeOChxZDwR/5xW6p1vq/wBbMsdM6fS5mK47hjzBc0H+f6ha3c2uln+h+hv/AO5N/pLr/q79WsXodJIPr5loHr5JEEx/g6m/4Olv7n/biu9L6TgdJxhjYNQrZy93L3u/ftsPuseriEp2KjpFGTLY4IDhh+Mv7"
        "ykkkkxhf//V9VSSSSUpJJJJSkkkklKWB1z6m9N6s92RWTh5jvpXVgFrz/3Yp9rbf+M/R3f8Kt9JOjxX6d1+Pj4v1d8Xg+bHoH3/cegOLsA2OqBmcN29h834N3539Wi3/jFNn1++sGGfSzaqXObyL630Wf2vds/8BXoyFk/z"
        "D/gpT/XEf2tqX+tGMnzqf/ReHH+M60N92BUT4jJgf+eChW/4y89+mPi4zD23WOt/6LBQtGn/AJUb/a/Iuowv5sofq+gH2yW3y/QR+ssn/evBftj699XEYtdzK7ODRSKWf2cnK/75eiYv1B63m2/aOp5DaHO+k97nZN/w3OP"
        "pt/7etXoSSPqr0cP03X/rK/Ve3/gfM43Sfql0Xpbm211G/JbqMi8h7wfGsQ2qn/rNbFspJKGXFfqu/FqZOPi9d3/WUkkkgsUkkkkp/9kAdjhCSU0EIQAAAAAAVQAAAAEBAAAADwBBAGQAbwBiAGUAIABQAGgAbwB0AG8Acw"
        "BoAG8AcAAAABMAQQBkAG8AYgBlACAAUABoAG8AdABvAHMAaABvAHAAIABDAFMAMwAAAAEAOEJJTQQGAAAAAAAHAAgAAAABAQAAIf8LSUNDUkdCRzEwMTL/AAAMSExpbm8CEAAAbW50clJHQiBYWVogB84AAgAJAAYAMQAAY"
        "WNzcE1TRlQAAAAASUVDIHNSR0IAAAAAAAAAAAAAAAEAAPbWAAEAAAAA0y1IUCAgAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAARY3BydAAAAVAAAAAzZGVzYwAAAYQAAABsd3RwdAAA"
        "AfAAAAAUYmtwdAAAAgQAAAAUclhZWgAAAhgAAAAUZ1hZWgAAAiwAAAAUYlhZWgAAAkAAAAAUZG1uZAAAAlQAAABwZG1kZAAAAsQAAACIdnVlZAAAA0wAAACGdmll/3cAAAPUAAAAJGx1bWkAAAP4AAAAFG1lYXMAAAQMAAA"
        "AJHRlY2gAAAQwAAAADHJUUkMAAAQ8AAAIDGdUUkMAAAQ8AAAIDGJUUkMAAAQ8AAAIDHRleHQAAAAAQ29weXJpZ2h0IChjKSAxOTk4IEhld2xldHQtUGFja2FyZCBDb21wYW55AABkZXNjAAAAAAAAABJzUkdCIElFQzYxOT"
        "Y2LTIuMQAAAAAAAAAAAAAAEnNSR0IgSUVDNjE5NjYtMi4xAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAABYWVogAAAAAAAA81EAAf8AAAABFsxYWVogAAAAAAAAAAAAAAAAAAAAA"
        "FhZWiAAAAAAAABvogAAOPUAAAOQWFlaIAAAAAAAAGKZAAC3hQAAGNpYWVogAAAAAAAAJKAAAA+EAAC2z2Rlc2MAAAAAAAAAFklFQyBodHRwOi8vd3d3LmllYy5jaAAAAAAAAAAAAAAAFklFQyBodHRwOi8vd3d3LmllYy5j"
        "aAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAABkZXNjAAAAAAAAAC5JRUMgNjE5NjYtMi4xIERlZmF1bHQgUkdCIGNvbG91ciBzcGFjZSAtIHNSR0L/AAAAAAAAAAAAAAAuSUVDIDYxOTY"
        "2LTIuMSBEZWZhdWx0IFJHQiBjb2xvdXIgc3BhY2UgLSBzUkdCAAAAAAAAAAAAAAAAAAAAAAAAAAAAAGRlc2MAAAAAAAAALFJlZmVyZW5jZSBWaWV3aW5nIENvbmRpdGlvbiBpbiBJRUM2MTk2Ni0yLjEAAAAAAAAAAAAAAC"
        "xSZWZlcmVuY2UgVmlld2luZyBDb25kaXRpb24gaW4gSUVDNjE5NjYtMi4xAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB2aWV3AAAAAAATpP4AFF8uABDPFAAD7cwABBMLAANcngAAAAFYWVog/wAAAAAATAlWAFAAAABXH"
        "+dtZWFzAAAAAAAAAAEAAAAAAAAAAAAAAAAAAAAAAAACjwAAAAJzaWcgAAAAAENSVCBjdXJ2AAAAAAAABAAAAAAFAAoADwAUABkAHgAjACgALQAyADcAOwBAAEUASgBPAFQAWQBeAGMAaABtAHIAdwB8AIEAhgCLAJAAlQCa"
        "AJ8ApACpAK4AsgC3ALwAwQDGAMsA0ADVANsA4ADlAOsA8AD2APsBAQEHAQ0BEwEZAR8BJQErATIBOAE+AUUBTAFSAVkBYAFnAW4BdQF8AYMBiwGSAZoBoQGpAbEBuQHBAckB0QHZAeEB6QHyAfoCAwIMAv8UAh0CJgIvAjg"
        "CQQJLAlQCXQJnAnECegKEAo4CmAKiAqwCtgLBAssC1QLgAusC9QMAAwsDFgMhAy0DOANDA08DWgNmA3IDfgOKA5YDogOuA7oDxwPTA+AD7AP5BAYEEwQgBC0EOwRIBFUEYwRxBH4EjASaBKgEtgTEBNME4QTwBP4FDQUcBS"
        "sFOgVJBVgFZwV3BYYFlgWmBbUFxQXVBeUF9gYGBhYGJwY3BkgGWQZqBnsGjAadBq8GwAbRBuMG9QcHBxkHKwc9B08HYQd0B4YHmQesB78H0gflB/gICwgfCDIIRghaCG4IggiWCKoIvgjSCOcI+wkQCSUJOglPCWT/CXkJj"
        "wmkCboJzwnlCfsKEQonCj0KVApqCoEKmAquCsUK3ArzCwsLIgs5C1ELaQuAC5gLsAvIC+EL+QwSDCoMQwxcDHUMjgynDMAM2QzzDQ0NJg1ADVoNdA2ODakNww3eDfgOEw4uDkkOZA5/DpsOtg7SDu4PCQ8lD0EPXg96D5YP"
        "sw/PD+wQCRAmEEMQYRB+EJsQuRDXEPURExExEU8RbRGMEaoRyRHoEgcSJhJFEmQShBKjEsMS4xMDEyMTQxNjE4MTpBPFE+UUBhQnFEkUahSLFK0UzhTwFRIVNBVWFXgVmxW9FeAWAxYmFkkWbBaPFrIW1hb6Fx0XQRdlF4k"
        "X/64X0hf3GBsYQBhlGIoYrxjVGPoZIBlFGWsZkRm3Gd0aBBoqGlEadxqeGsUa7BsUGzsbYxuKG7Ib2hwCHCocUhx7HKMczBz1HR4dRx1wHZkdwx3sHhYeQB5qHpQevh7pHxMfPh9pH5Qfvx/qIBUgQSBsIJggxCDwIRwhSC"
        "F1IaEhziH7IiciVSKCIq8i3SMKIzgjZiOUI8Ij8CQfJE0kfCSrJNolCSU4JWgllyXHJfcmJyZXJocmtyboJxgnSSd6J6sn3CgNKD8ocSiiKNQpBik4KWspnSnQKgIqNSpoKpsqzysCKzYraSudK9EsBSw5LG4soizXLQwtQ"
        "S12Last4f8uFi5MLoIuty7uLyQvWi+RL8cv/jA1MGwwpDDbMRIxSjGCMbox8jIqMmMymzLUMw0zRjN/M7gz8TQrNGU0njTYNRM1TTWHNcI1/TY3NnI2rjbpNyQ3YDecN9c4FDhQOIw4yDkFOUI5fzm8Ofk6Njp0OrI67zst"
        "O2s7qjvoPCc8ZTykPOM9Ij1hPaE94D4gPmA+oD7gPyE/YT+iP+JAI0BkQKZA50EpQWpBrEHuQjBCckK1QvdDOkN9Q8BEA0RHRIpEzkUSRVVFmkXeRiJGZ0arRvBHNUd7R8BIBUhLSJFI10kdSWNJqUnwSjdKfUrESwxLU0u"
        "aS+JMKkxyTLpNAk3/Sk2TTdxOJU5uTrdPAE9JT5NP3VAnUHFQu1EGUVBRm1HmUjFSfFLHUxNTX1OqU/ZUQlSPVNtVKFV1VcJWD1ZcVqlW91dEV5JX4FgvWH1Yy1kaWWlZuFoHWlZaplr1W0VblVvlXDVchlzWXSddeF3JXh"
        "pebF69Xw9fYV+zYAVgV2CqYPxhT2GiYfViSWKcYvBjQ2OXY+tkQGSUZOllPWWSZedmPWaSZuhnPWeTZ+loP2iWaOxpQ2maafFqSGqfavdrT2una/9sV2yvbQhtYG25bhJua27Ebx5veG/RcCtwhnDgcTpxlXHwcktypnMBc"
        "11zuHQUdHB0zHUodYV14XY+/3abdvh3VnezeBF4bnjMeSp5iXnnekZ6pXsEe2N7wnwhfIF84X1BfaF+AX5ifsJ/I3+Ef+WAR4CogQqBa4HNgjCCkoL0g1eDuoQdhICE44VHhauGDoZyhteHO4efiASIaYjOiTOJmYn+imSK"
        "yoswi5aL/IxjjMqNMY2Yjf+OZo7OjzaPnpAGkG6Q1pE/kaiSEZJ6kuOTTZO2lCCUipT0lV+VyZY0lp+XCpd1l+CYTJi4mSSZkJn8mmia1ZtCm6+cHJyJnPedZJ3SnkCerp8dn4uf+qBpoNihR6G2oiailqMGo3aj5qRWpMe"
        "lOKWpphqmi6b9p26n4KhSqMSpN6mpqv8cqo+rAqt1q+msXKzQrUStuK4trqGvFq+LsACwdbDqsWCx1rJLssKzOLOutCW0nLUTtYq2AbZ5tvC3aLfguFm40blKucK6O7q1uy67p7whvJu9Fb2Pvgq+hL7/v3q/9cBwwOzBZ8"
        "Hjwl/C28NYw9TEUcTOxUvFyMZGxsPHQce/yD3IvMk6ybnKOMq3yzbLtsw1zLXNNc21zjbOts83z7jQOdC60TzRvtI/0sHTRNPG1EnUy9VO1dHWVdbY11zX4Nhk2OjZbNnx2nba+9uA3AXcit0Q3ZbeHN6i3ynfr+A24L3hR"
        "OHM4lPi2+Nj4+vkc+T85YTmDeaW5x/nqegy6LxU6Ubp0Opb6uXrcOv77IbtEe2c7ijutO9A78zwWPDl8XLx//KM8xnzp/Q09ML1UPXe9m32+/eK+Bn4qPk4+cf6V/rn+3f8B/yY/Sn9uv5L/tz/bf//ACwAAAAAaQBGAAAI"
        "/gA3CRxIsKDBgwgTKlzI0NCcN3cMLmJIsaLFixYH3ZnTBkyXK06C8JiR4gUNH1fWuDE0cCLGlzBjDrzz5gyYm2fauIFTp87GOW6+XPERA8UIGk3WsJTJtKnCOm3OrJkzyOIiN1dqjBhRo0obgS6dipX56GDZsgLRnkUrcM2"
        "PEyNSMPGzKezYuxgfsbWod+AaHihQ1HCzaS/ew4gJstUC4wQMM4kjS04r8M0NFCm8FJ7MGW/ZQT4wg9nc+e6cyGh9OP5quPTLsmBAZCHtua4OwXpou8676Q2IDx/eRGZZZwYJIbubPiJSYYtuxF8Cf0nedE3nskJG6LBbm7"
        "fC1onh/sAgoeU59bGMli58pB5hExI32iM0JF8h3b4VGW0a9AchI/0JAbiJIf/pV2B69RlUlhspoDCdQhogsEACZay3SRUmyPCCDXWYN+AmdDygAAIO0HFQHAwokAAaHxqEhgYOLLAAAwsoIOONMj6wQRwL+TACcgexhEUAR"
        "AJwQYsK/TBBCBXIUBV4mzxApABibCKglZuwIYAAAVR4YgEAhCnmmGQCkEAeCJWlxVERISQlAgxwKQWSCQ1BQQkf5EBZQRkEsGWVCcZRwJYVXsmSEQAUAAEGF2DA6KOONgrBAAAoQSdlb6iQwoMEsQQFkS2UQaQDDKEVRAUm"
        "VKBDQSyxAMAA/gLMmWCWgwrgZaebaOAniwuhMUAAHVw60CCXOYEQAwEUIMcmFxBp6awF2WnCBz8MxJISAfwKhbADsQGrrcL2SeWH9FnJEiP9hQGAAEdgiRAQI/jA6iZJEMmCQGJsmQC0Bg3xQQkVGFvVkFwmwS2WghJqkH7"
        "iknGwQHEosWxCLD0xAg4G5YGAAAeYyFKfARTxsGL+AkzFJmlsGfLItCpcn7hVOqUmCjTg2gKRIhsyySZlbFmAiTt/J5CdInyQBSQYAOCBhwOVlTC47bGkK7gyb7IFCjAUJMegCOxhba5EcsAyaUSAUHQWhlhaURzf3jrQzk"
        "Vku8ADDThgdwN13+0A/gQb8KoQF1h/zQGRSLS4c8pbprFJ0BTZ8JsLjyTCMiQtQ22QGlwSqfnmnL+qeIIzZz2QGr8m4DVLk0wCyCYbEJnB4qVuQgUIIEAuhQBYjM22wgaxFIYFCRxQwAEGDG/A8ccLD2awoG9yxQkxCPSxn"
        "9seZEetYWyCiNBafACCCXiM4acAU7DMRq0s1hf0HnS07/77dOQhxq9HNs/ECTV8mO8AA1hwwf8A/B8GDsAlCMAuTZvogvdAcIZNyKEB45OVQrzFO6aQTgAPWAhgdjAQC/gpAGUak59+FQVuTSQ2IqhAF/azCTskYEvk45Z+"
        "nlah9gSNDWronUHS8KsMJkQP/jRAgRBYIgU/ISACEEgiBJCoxAhEYAFbWgDjCDKRNdDuA0/YjNMIuCUJHuR8FZTeJqbgAAFsYEAsIdCA9FOulAXAhwhxgwpQcAWBOMBPc1oIHQxQMDqV5Q0lSGG12MISHsIwd/0xCA3nhYQ"
        "wrYAibHsjAgGXAsh8SgAW+BDlWrSUQ/npAGhq2ibqYILf6EmH+TqksMBoOTFGgUsaYEMZ0EDLMsyyDGlAQxnUEIVfYWA+myDCCWYQESgGIHf8EkgeCBiAe7GkLHdwAe1c8CQd3k4AsELmQCh3Pli5bXRb+hYMx4nNcIIpZw"
        "cBIgrk5YEwZXKKwGwkkViklzq4QAL//oID0/TzKRACoHoEYUORHBa1TdyMcwgFYZFKlBAwpOAE5WFBB/o2NoHwYQUTlWAXguAEIlgHSmKEwgZWwAEPhHIgdMBoB3iEkDIcwQMd6ABMY0pTmrYgCqhDSBAwIxyxtAak57HIH"
        "IrCA4FMolwYIZAh9jKRRQBVMen5T5CQ6h+llotACFojVQ9ShcCUhztBDatB7kCDFMwgN2JNq0ScF5gqqPWtBblDDFQAgw7B9a5MaGtF7yqZKqLgBcTk61sXcYMGOeepMFlEMgW7GSdgpqiR0cNpGJuQiUQHBiroaWLu8BV4"
        "UnYg4nkBROsSWch89iCBqMELUACEzmQhEKc1/ggPUqACGqAVNRdqoGfv6pKdYpYwYL1LWbxAhNiiRZhFmU5wD1OHGHyBMKcVAgqKUh7EjkUHQdBCVfg6EUOEZjx1dM1waVCF6Vi3Mzurw20aU8fz3kUHN9iBezuzBpI0pjx"
        "Mk4wXAmMFtbpEC7RdLaeCGpoXrHC5pakDEEgAA6w1MK1zgMELUnDg3YSlC0WJwQl0cBoEo3cT+22wc8RbGdW8YI5ZhCtLrAA9FPQ3v05Zyh2eMMcG1+DBgnWCY9Zp1/le5A5WaAwMGvSEqng4rE/YcQxWuCeZ2AUOT9CwhE"
        "mAgwcfWaxXS0GDdfAF9eCnIk4diB6+EAQVOKYxM9DCvUR8fJ411ABrDbrBFfQpkb488yB3+AITgqjlh87gCm2KLatojLU5pmAHTzDDHJ66hzmAoQo8KMpfG4wCHGgBrVe+K1rgQITV0rVBWKuBD5pgBS2YWgtXeIJIyhqYk"
        "lAaBkIwg0syzVi0zMEKlzmBq1MA6hOgwNeBwUyfW12DIHgh0IKmSNAWsYah0EBTwu4zr3mNGRjo4AldcIPXWpJsqxAkD2voghWI4IMd6AAITXBCFa7ghTNMtiVsvkhAAAA7")

        # This is the main/root widget 
        self.parent = parent                                                                        # This is the parent widget 
        self.parent.configure(background='white')                                                   # Change background color to white

        # Create a frame for the header
        self.header_string = "AXIS' Mail Label Application"                                         # Create a string to insert into lable
        self.header_frame = tk.Frame(self.parent,padx=5,pady=5,highlightbackground="black", highlightcolor="black", highlightthickness=1) # Create a frame and tie it to the parent widget
        self.header_frame.pack(side='top',fill=BOTH,expand = True)                                  # Set the frame to stick to the top of parent widget. This also stretches the widget across the x and y lengths
        #self.logo = tk.PhotoImage(file="axis_logo_icon.gif")                                       # Set AXIS' logo - NOTE: If we do this method we have to add file using --add data or modifying the spec file, then adding the tk temp dir to path 
        self.logo = tk.PhotoImage(data=self.axis_logo_base64_string)                                # Set AXIS' logo - Just add the raw base64 string       
        self.label= tk.Label(self.header_frame,image=self.logo,justify=LEFT)                        # Associate logo with label
        self.label.pack(side='top')                                                                 # Pack the lable to the top
        self.label= tk.Label(self.header_frame,text=self.header_string,justify=LEFT,font="-weight bold") # Create a label widget and tie it to the header frame w/ label string
        self.label.pack(side='top')                                                                 # Pack the lable to the top (it will be below the element that has attribute slide='top' first)
        self.header_frame.configure(background='white')                                             # Change background color to white
        self.label.configure(background='white')                                                    # Change background color to white        

        # Create a frame for description
        self.description_string = """This program accepts an excel spreadsheet containing borrower addresses and automatically generates a formatted word document that can be used to print mailing lables. \nFollow the steps below to complete the process.\n"""  # Create a string to insert into lable
        self.description = tk.Frame(self.parent,padx=5,pady=5,highlightbackground="black", highlightcolor="black", highlightthickness=1) # Create a frame and tie it to the parent widget
        self.description.pack(side='top',fill=BOTH,expand = True)                                   # Associate it with the top of parent widget (it will be right below any widget above it that is also associated with side='top' to parent). This also stretches the widget across the x and y lengths
        self.label= tk.Label(self.description,text=self.description_string,justify=LEFT)            # Create a label widget and tie it to the description frame w/ label string
        self.label.pack(side=LEFT)                                                                  # Pack the lable to the left
        self.description.configure(background='white')                                              # Change background color to white
        self.label.configure(background='white')                                                    # Change background color to white

        # Create a frame for main content 
        self.content_box = tk.Frame(self.parent,padx=5,pady=5,highlightbackground="black", highlightcolor="black", highlightthickness=1) # Create a frame and tie it to the parent widget
        self.content_box.pack(side='top',fill=BOTH,expand = True)                                   # Associate it with the top of parent widget (it will be right below any widget above it that is also associated with side='top' to parent). This also stretches the widget across the x and y lengths
        self.content_box.configure(background='white')                                              # Change background color to white

        # Create a frame for the "step 1" section / downloading sample document
        self.step1_title_string = "Step 1\n"                                                        # Create a title string to insert into lable
        self.step1_instructions_string = "Download this sample template, enter borrower addresses,\nthen upload filled out document in 'Step 2'\n" # Create an instruction string to insert into lable
        self.step1_frame = tk.Frame(self.content_box,padx=5,pady=5)                                 # Create a frame and tie it to the content_box widget
        self.step1_frame.pack(side='left')                                                          # Set the frame to stick to the left of content_box widget
        self.label= tk.Label(self.step1_frame,text=self.step1_title_string,justify=LEFT)            # Create a label widget and tie it to the step1 frame w/ label string
        self.label.pack(side='top')                                                                 # Pack the lable to the top
        self.label.configure(background='white')                                                    # Change background color to white
        self.label= tk.Label(self.step1_frame,text=self.step1_instructions_string,justify=LEFT)     # Create a label widget and tie it to the step1 frame w/ label string
        self.label.pack(side='top')                                                                 # Pack the lable to the top (will be right below the other label since that was side='top' first)
        self.button = tk.Button(self.step1_frame, text='Download Sample Document', command=self.export_template) # Create a button and tie it to the step1 frame widget
        self.button.pack(side='top')                                                                # Pack the button to the top (will be right below the other two labels since that was side='top' first)
        self.step1_frame.configure(background='white')                                              # Change background color to white
        self.label.configure(background='white')                                                    # Change background color to white

        # Create a frame for the "step 2" section / downloading sample document
        self.step2_title_string = "Step 2\n"                                                        # Create a title string to insert into lable
        self.step2_instructions_string = "Upload spreadsheet with borrower email addresses.\nOnce completed, export mailing lables in 'Step 3'\n" # Create an instruction string to insert into lable
        self.step2_frame = tk.Frame(self.content_box,padx=5,pady=5)                                 # Create a frame and tie it to the content_box widget
        self.step2_frame.pack(side='left')                                                          # Set the frame to stick to the left of content_box widget (it will be to the right of any other widget set to left before it, and any widget set to left after it will be on its right)
        self.label= tk.Label(self.step2_frame,text=self.step2_title_string,justify=LEFT)            # Create a label widget and tie it to the step1 frame w/ label string
        self.label.pack(side='top')                                                                 # Pack the lable to the top
        self.label.configure(background='white')                                                    # Change background color to white
        self.label= tk.Label(self.step2_frame,text=self.step2_instructions_string,justify=LEFT)     # Create a label widget and tie it to the step1 frame w/ label string
        self.label.pack(side='top')                                                                 # Pack the lable to the top (will be right below the other label since that was side='top' first)
        self.button = tk.Button(self.step2_frame, text='Upload Address Document', command=self.upload)   # Create a button and tie it to the step2 frame widget
        self.button.pack(side='top')                                                                # Pack the button to the top (will be right below the other two labels since that was side='top' first)
        self.step2_frame.configure(background='white')                                              # Change background color to white
        self.label.configure(background='white')                                                    # Change background color to white

        # Create a frame for the "step 3" section / downloading sample document
        self.step3_title_string = "Step 3\n"                                                        # Create a title string to insert into lable
        self.step3_instructions_string = "Export mailing lables and save to file system\nbelow.\n"  # Create an instruction string to insert into lable
        self.step3_frame = tk.Frame(self.content_box,padx=5,pady=5)                                 # Create a frame and tie it to the content_box widget
        self.step3_frame.pack(side='left')                                                          # Set the frame to stick to the left of content_box widget (it will be to the right of any other widget set to left before it, and any widget set to left after it will be on its right)
        self.label= tk.Label(self.step3_frame,text=self.step3_title_string,justify=LEFT)            # Create a label widget and tie it to the step1 frame w/ label string
        self.label.pack(side='top')                                                                 # Pack the lable to the top
        self.label.configure(background='white')                                                    # Change background color to white
        self.label= tk.Label(self.step3_frame,text=self.step3_instructions_string,justify=LEFT)     # Create a label widget and tie it to the step1 frame w/ label string
        self.label.pack(side='top')                                                                 # Pack the lable to the top (will be right below the other label since that was side='top' first)
        self.button = tk.Button(self.step3_frame, text='Save Formatted Results', command=self.export_results) # Create a button and tie it to the step3 frame widget
        self.button.pack(side='top')                                                                # Pack the button to the top (will be right below the other two labels since that was side='top' first)
        self.step3_frame.configure(background='white')                                              # Change background color to white
        self.label.configure(background='white')                                                    # Change background color to white

        # Create a frame for the footer
        self.footer_string = "Contact developer at TFye@axis-amc.com if there are errors"           # Create a string to insert into lable
        self.footer_frame = tk.Frame(self.parent,padx=5,pady=5,highlightbackground="black", highlightcolor="black", highlightthickness=1) # Create a frame and tie it to the parent widget
        self.footer_frame.pack(side='top',fill=BOTH,expand = True)                                  # Set the frame to stick to the top of parent widget. Stretch widget across x and y lengths
        self.label= tk.Label(self.footer_frame,text=self.footer_string,justify=LEFT)                # Create a label widget and tie it to the header frame w/ label string
        self.label.pack(side=LEFT)                                                                  # Pack the lable to the left
        self.footer_frame.configure(background='white')                                             # Change background color to white
        self.label.configure(background='white')                                                    # Change background color to white

        # Testing Block
        # Uncomment the two bottom lines for debugging to create a text area in application in class vars. Then, include the code expert in quotes in any method necessary. 
        #    Anything can be printed. In the example to the left the 'self.data' member would be printed in any method it is added to => "self.text.insert('end', str(self.data) + '\n')""                 
        # self.text = tk.Text(self.parent)
        # self.text.pack(side='top')

    # Function to display the datafrom excel spreadsheet
    def export_template(self):                                            
        saveFile = asksaveasfile(mode = 'w',defaultextension=".xls")  # Prompt to allow user to save output file to their desired location
        if saveFile is None:                                           # If no file was created...                                                                                        
            messagebox.showinfo("Error","Application Error")           # Alert user there was an error 
        rawFilePathString = str(saveFile)                              # Store the raw file string into a local variable. It will be in format '<_io.TextIOWrapper name='C:/Dir/File.docx' mode='w' encoding='cp1252'>'
        start = 'name=\''                                              # We want to find substring to the left of this 
        end = '\' mode'                                                # We want to find substring to the right of this 
        self.sampleFilePath = rawFilePathString[rawFilePathString.find(start)+len(start):rawFilePathString.rfind(end)] # Store the path of the output file to our class variable 
        saveFile.close()                                               # Close the file so the algorithm can work with it          
        algorithm = Algorithm()                                        # Create a new algorithm object 
        spreadsheetArray = [ ["Borrower Name","Borrower Subject Street","Borrower City, State, Zip" ] ] # Provide the rows / columns to be generated in template 
        algorithm.generate_template(spreadsheetArray, self.sampleFilePath)    # Run the custom algorithm to create the xls export of file

    # Function to load the excel spreadsheet from user's desktop
    def upload(self):
        name = askopenfilename(filetypes=[('Excel', ('*.xls', '*.xlsx'))]) # Prompt to request user to upload a file for evaluation
        excel = Excel()                                                    # Create a new excel object 
        self.data = excel.read_data_from_excel(name)                       # Call the function to read spreadsheet data into an array of arrays 

    # Function to display the datafrom excel spreadsheet
    def export_results(self):                                            
        if self.data is None:                                              # ask for file if not loaded yet
            messagebox.showinfo("Error","Upload 'Address Document' File First") # Prompt the user to upload a file first
        if self.data is not None:                                          # display if loaded
            saveFile = asksaveasfile(mode = 'w',defaultextension=".docx")  # Prompt to allow user to save output file to their desired location
            if saveFile is None:                                           # If no file was created...                                                                                        
                messagebox.showinfo("Error","Application Error")           # Alert user there was an error 
            rawFilePathString = str(saveFile)                              # Store the raw file string into a local variable. It will be in format '<_io.TextIOWrapper name='C:/Dir/File.docx' mode='w' encoding='cp1252'>'
            start = 'name=\''                                              # We want to find substring to the left of this 
            end = '\' mode'                                                # We want to find substring to the right of this 
            self.outputFilePath = rawFilePathString[rawFilePathString.find(start)+len(start):rawFilePathString.rfind(end)] # Store the path of the output file to our class variable 
            saveFile.close()                                               # Close the file so the algorithm can work with it       
            algorithm = Algorithm()                                        # Create a new algorithm object 
            algorithm.generate_results(self.data, self.outputFilePath)     # Run the custom algorithm to create the mailing list
            # Reset all local vars so user can upload a new file without having to close the program
            self.filename = None
            self.outputFilePath = None
            self.sampleFilePath = None        
            self.data = None   

    # Function to extract dependencies / data files built into tkinter's exe (doesn't work as written)
"""    def resource_path(relative_path):
         # Get absolute path to resource, works for dev and for PyInstaller
        try:
            # PyInstaller creates a temp folder and stores path in _MEIPASS
            base_path = sys._MEIPASS
        except Exception:
            base_path = os.path.abspath(".")

        (return os.path.join(base_path, relative_path)
 """
class Excel:

    # This function is responsible for opening excel file, reading data from it, and returning an array or arrays that hold all spreadsheet data. 
    def read_data_from_excel(self,filepath):

        # File name to be opened (must be saved in same directory as .py exe's or full path provided)
        inputFile = filepath

        # Read data from an excel file  
        wb = xlrd.open_workbook(inputFile)                        # Open the file in question
        for s in wb.sheets():                                     # For all sheets that exist in the workbook_rb 
            values = []                                           # Initialize an array that will hold all values in the workbook_rb   
            for row in range(s.nrows):                            # For each row on the current sheet in workbook_rb
                col_value = []                                    # Initialize an array that will hold all column values for specified row 
                for col in range(s.ncols):                        # For each column that is in the current row 
                    value  = (s.cell(row,col).value)              # Find the value of the cell, the col reference will increment each loop 
                    try : value = str(int(value))                 # See if we can cast the value to an in within a string 
                    except : pass                                 # Include exception 
                    col_value.append(value)                       # Append the cell value to the column value array 
                values.append(col_value)                          # Append the column value array 

        # return the multi-dimentional data array to calling function 
        return values

class Algorithm:

    def generate_template(self,spreadsheetArray,outputFilePath):

        # Before we can modify the file created above, the directory the user saved the file in needs to be added to path (py only executes in dir it was run in)
        # Variable declarations
        length = len(outputFilePath)     # Find the length of the string
        file_path = None                 # Intialize a variable_name
        file_name = None                 # Initialize a variable for file_name
        leftMostBackSlashIndex = None    # Initialize a variable to hold the index of the left most "/"

        # Iterate through the full file path and find the index of the last "/" in the text string. For example: 
        #    "C:/Directory1/Directory2/file.docx"
        #                             ^
        #                             With this index we can split the string into two components:
        #                             1. The path to directory
        #                             2. The file name 
        # NOTE: This will only work in Windows environment. Linux uses '\' as path. Code would need to be updated to use in both env.
        while length != 0:                                      # While length is greater than zero 
            if outputFilePath[length-1] != "/":                 # If the character at the current index in text string is not "/"...
                length -= 1                                     # Decrement our length by one to check the next character
            else:                                               # If we did land on the first "\" character reading from left to right... 
                leftMostBackSlashIndex = length                 # Set the index of our left most back slash
                break

        # Now splice the string into two parts
        if leftMostBackSlashIndex == None:                                          # If we haven't found the "/" and intialized our index var... 
                messagebox.showinfo("Error","Application Error")                    # Alert user there was an error     
        else:                                                                       # Otherwise...
            file_path = outputFilePath[0:leftMostBackSlashIndex]                    # Splice first half of string to obtain the path 
            file_name = outputFilePath[leftMostBackSlashIndex:len(outputFilePath)]  # Splice the second half of string to obtain the file name

        # Add tkinter program's document into the same scope/file path as location the user choose to save file, tried to use sys.path.append(file_path) & os.path.join(file_path, file_name) without success
        os.chdir(file_path)                           # Change directory to dir the user choose to save their doc in 

        # XLS: create a new temp file for writing 
        workbook = xlwt.Workbook(encoding = 'ascii')   # Create a new workbook_rb
        worksheet = workbook.add_sheet('Sheet 1')      # Add a sheet to the workbook_rb

        # Set the payload equal to spreadSheetArray parameter
        payload = spreadsheetArray

        # Write data from array into spreadsheet
        rowIndex = 0                                            # Initialize row index to zero, this will be incremented as we step through sub-arrays in outer array
        for row in payload:                                     # For each row in the payload array (aka outer array)...               
            columnIndex = 0                                     # Initialize a column index to zero, this will be incremented as we step through current sub-array (aka current row)                    
            for cell in row:                                    # For each item in the row array (aka cell)...
                worksheet.write(rowIndex, columnIndex, cell)    # Write data relative the appropriate row and column index
                columnIndex += 1                                # Increment the column index so we move to next column and prepare to write to next cell during next pass
            rowIndex += 1                                       # Increment the row index so we move to next line as we prepare to write next row

        # Save file for persistence of changes
        workbook.save(file_name)                                # Re-write the read only version of excel sheet by saving the copy in the same directory with the same name

    def generate_results(self,spreadsheetArray,outputFilePath):

        # Before we can modify the file created above, the directory the user saved the file in needs to be added to path (py only executes in dir it was run in)
        # Variable declarations
        length = len(outputFilePath)     # Find the length of the string
        file_path = None                 # Intialize a variable_name
        file_name = None                 # Initialize a variable for file_name
        leftMostBackSlashIndex = None    # Initialize a variable to hold the index of the left most "/"

        # Iterate through the full file path and find the index of the last "/" in the text string. For example: 
        #    "C:/Directory1/Directory2/file.docx"
        #                             ^
        #                             With this index we can split the string into two components:
        #                             1. The path to directory
        #                             2. The file name 
        # NOTE: This will only work in Windows environment. Linux uses '\' as path. Code would need to be updated to use in both env.
        while length != 0:                                      # While length is greater than zero 
            if outputFilePath[length-1] != "/":                 # If the character at the current index in text string is not "/"...
                length -= 1                                     # Decrement our length by one to check the next character
            else:                                               # If we did land on the first "\" character reading from left to right... 
                leftMostBackSlashIndex = length                 # Set the index of our left most back slash
                break

        # Now splice the string into two parts
        if leftMostBackSlashIndex == None:                                          # If we haven't found the "/" and intialized our index var... 
            messagebox.showinfo("Error","Application Error")                    # Alert user there was an error     
        else:                                                                       # Otherwise...
            file_path = outputFilePath[0:leftMostBackSlashIndex]                    # Splice first half of string to obtain the path 
            file_name = outputFilePath[leftMostBackSlashIndex:len(outputFilePath)]  # Splice the second half of string to obtain the file name

        # Add tkinter into the same scope/file path as location the user choose to save file, tried to use sys.path.append(file_path) & os.path.join(file_path, file_name) without success
        os.chdir(file_path)                           # Change directory to dir the user choose to save their doc in 

        # Variable Declarations
        borrowerAddressArray = spreadsheetArray       # Define a local instance of our spreadsheetArray
        document = Document()                         
        style = document.styles['Normal']             # Access global style object
        font = style.font                             # Access font object
        font.name = 'Arial'                           # Change font 
        font.size = Pt(16)                            # Change font size 
        axis_mailhead_base64_string = ("iVBORw0KGgoAAAANSUhEUgAAAX8AAACaCAYAAABBukjiAAAAAXNSR0IArs4c6QAAAARnQU1BAACxjwv8YQUAAAAJcEhZcwAADsMAAA7DAcdvqGQAAC3zSURBVHhe7Z0JnFTFufaT35fl5ub"
        "+bq7JZ2K+xF2TmMWoMVezoSYx7poYo5KYaGKiIogkxg13EJRNVtllUxBZBWUTBNnXYdh3GGAYZoBhmX2f56unznmna4rTPT0wMzL2+9eaPqeWt96qU/VU9enTzaegKIqipBwq/oqiKCmIir+iKEoKouKvKIqSgqj4K4qipCAq/oqiKC"
        "mIir+iKEoKouKvKIqSgqj4K4qipCAq/oqiKCmIir+iKEoKouKvKIqSgqj4K4qipCAq/oqiKCmIir+iKEoKouKvKIqSgqj4K4qipCAq/oqiKCmIir+iKEoKouKvKIqSgqj4K4qipCAq/oqiKCmIir+iKEoKouKvKIqSgqj4K4qipCAq/"
        "oqiKCmIir+iKEoKouKvKIqSgqj4K4qipCAq/oqiKCmIir+iKEoKouKvKIqSgqj4K4qipCAq/oqiKCmIir+iKEoKouKvKHVQXV1tQw085LkN/L8aVUFk+DcWVxUe26PqyppUhCnWhj0KXhWlqVDxV5SGgNpdaf4wRGKWguoqVNUsJMHS"
        "UGmCXTjChURRmgoVf0WpC+oyN+4VRqzLKlBeWo7SwlIUHCnEkewjyNmVg8yNe7B1yQas/zANm+evRcaqHcjZlo38nDyUF5QfL+xW7I38M9gloNJkUfVXmg4Vf0WJA6W4srIKpSXlKDFiX5JfjMJjBSjKK0LR4ULk7T+KQ7sPYv+WLOx"
        "evQub5q9H+rSVWDlpGVZOWIqV4xdj5biF5nwRts3fiNztOSjLK3UWAi4ARvSrK8wiENw4UpSmQsVfUeLAzXm52ekXFRrRzy9EcX6RDaVFxSgvLEGZWQxK84pRcszE5xah8GABjmUexYGtB7AnLQOb56zF8jHzMLvvZMzsMQEfDZiO9M"
        "nLsS99L0qOFAcVmB1/sACo9CtNi4q/osSBglxZWYmK8nL7arbnYQo/rJXAuDjCbaIr8suRuy0Ha6cux9Qu72DUPwdi0oujsWbyCuRnHkZ1VYXJxs8BwjKK0kSo+CtKiDzVIyGKIL52qHmmh7du7O2bIPBvjRVzUJhdgHXvp+Gdp4fjj"
        "Yf6YN7r03Fkx0Gz+Tc5q/ihQphXUZoAFX+leSLaezJQyPlJbmjIFXx7HJ5SwmMp9ccuJuFxVVkVcrfkYMGgDzCs1et4r9NYHNucY94ihEsI6zULQbCInEytipIYFX8ldQnFP9ilB4JrH9UsKTdiHOzEZUE4GRm24m8XgNCKMZ2/7xjS"
        "xi/HqHYD8F6H0Tiy7WCwDhmq+RkA/TmpWhUlMSr+SgpDeQ0fsqTQl1Rh17JtWDBqNvak70RlaUWYrz5QsGMhWDycwJeQgpx8rHpnKYY/3A/zB0xDwZ4jNj24YSRfCFOUxkHFX/nYqNkRhyEZavLzP69MlI3EeQI7lFnzgqr8KswfPhc"
        "DH+yNNe+vQEWxeQdg0463Gx/mdYO9mVMTaKrafnDMYBaArHzMHzwHI1r1xvpJy4KngAz2JpDNbE8VpcFR8Vc+NqyIOyEZEpWJsuHnq5XHpsXkuaqoCutmpGNWn/ewZ9UOVJfZZcEkHG83PswbHfifvZnD+ztVoW3DkR2HMfbJ4Zjw7E"
        "hkrDD1hknW1/BWkKI0NCr+SvNCtDRJohYEYuNDWxTl4GcWjDSXVaC6KHbP3wo/Xx3i2awTU4412SXACrutwP6/7r10DHm4D+YPnYXCnLwwO304wboUpQ5U/JWPDQpgRUUFyssrUFlhRDHcYScUV5NWZfJWGZGuKjevfEzSxMl/Lq6dK"
        "mM7qKvcvlbJo5VhFpatMlvuiqoKVFQafyodf8LgQtv2OwDGVjLB1sdajKDb3b81Yo7DbX5xdhHe7TgGk54bhYzF22ycvfVzXM2K0jCo+CtNiggy/+7ffwATJryHAf1HYObMeTh4MDdIS7DjrTQ78kO7DmLl5CVY+uYCbJ23BQUHCoJE"
        "FuFCYMpbG1Y8gXJTZsvWXRg2bDQGDhyBtNVrUVJSatMCgrqKC0swZ/YCDBk6GosWrUB+XmA3cCXmD8+3bduFKe/OwNChbxmbIzHAhlHh64jwdSQGDRqFMWMmYMXydBQWlARmTLALVuintWzWgPQJyzD2sTewdOQ8VOaVMTaoTFEaARV"
        "/pcmgjon4V5jd/vDhY3DFlb/GV75yPn71q99i/PgpNs1ixTuERUIN5A7+2J6jeL/zO+h27fPoecvLWDx0HsoOBWLO3XLw/E6sfFr6evz1/rY486zv409/etCI/3r7biMgJq4HDx5Gm9aP47zzL8Xzz7+CzMysMIXE7K1btxmtW/8b3/"
        "3uj/Gtb11mwuVh+BEuDENwbtIvvBzfuejHpn23YtKk6SgoKLQ2AvGP1U2y1+3DxOfeMm2bgNxtB4LI2lkUpcFQ8VeakEDJKOBbt2bgllta4ocXt8APfvBznH/eJbj33lbYlZFp8wRiGyofX+TQCGZFcRn2LNqKcQ8PRYeLH0Pv33TCi"
        "uELUX4sWAB4/17YvHkXHnnkKbPAnI3LLr8KY96ejCNHeU/9eFXNzjmIO++8F5/+P6fh4Tb/RkbGnjCFeWM2+Q7ikh/+DN/73pVo0+ZJPPfcK2j/dAc8bUNHE1624ZlnOuOJJzrg9tvvxRf/8ww88MBj2L49I7RyPCUHizCr51SMf2YU"
        "di7aEsYqSuOg4q80OXn5BejffwTOPPNiPNr2OXvb54br78BF37nC3iaptL+JHxNce2b+2I2y/VMJlFUhc9EuTGg9Cp0ufRKv3/wq1r6zAuV55TWyvnNXphHhzjj3nB/i4ot/akR7DI4dC28RRUDxv/vuv+Izn/0K2jzyBDJ27w1TaDE"
        "m/h1f7oHzzGJ1551/w8qV65C9/yAy92aFYb8Ne03IyjqAnTv3YuzYd3HZZVeh3aPP2NtF8eD3DBaPmIdxT47A2vdWhrGK0jio+CtNCnf9q9PX43oj9t/77s8wY/p8HM7NR7eufXHeuZfghhvutLv1gEDxA+kNbujYc+c++d7FuzGuzU"
        "i8fNkTGPS77tg0ZT2qi6tw5EgeOr/aCxdc+CN8/3s/xeuvD0NBfhFLWfxbLiQ7+yD+2PJ+fP5zp6Nt26cc8Scx8X/hxS644IJL0arVYzhs6qmLzH378YZ5tzB37mLrV1yMS6snLsOEp0Zh5dhFYaSiNA4q/koTEYgt76v36DEAZ5xxI"
        "Vq3ftLsknNs/NKlabj55rvx1dMvMOmDUFrKDzyD++KB2Mt9fGPHvIQP4tjjzKUZZgEYYRaAxzGsZV/sm78XKz9ajVtuudPs+H+CLl36mh1/cK9d/IgRO8/OPoSWLf+Oz33ua3jkkSeRkXG8+DN3hw5djfhfgj//5SH7QXJxcQny8wvj"
        "hry8AhSXhB/g1sH6aWmY9MxbWDZ6/vGuKkoDouKvNCnz5y9DixY346LvXoGZs+aGIg8jzvno02cIzvzm93HFlddj3bqt4D+kEsDXmPjyX7+y34SSb0OZlwOr92NSu7fw4sX/xIh7BmD+yPno99ogjHl7ArL2BQuMLX2coMYisvfnGvH"
        "/hxH/M+znBLF7/kR8AXr3HoTvfPtyXHnFtejWbQCGDx+PIUPGYOhQJ5jzIYNNGDIao0dPwOKFK3Dk8DH7zicRG2etxeQXxhjxnxe4lji7opwwKv5KI1Jbvfj0TPtnOuL/ff0itPvn8zhsxNBlzdqNuOuuf+C//vssvPhSVxw+cjRMCR"
        "6JtPprbwOZdwHV5WYRqELwa/gmqbgSGTO3od8vX8EL338MM1+bhrRF6dhr6iwrcx6bDDU8KBUciY+B+D8Qiv/T3s6f70CCoyVLVuD22+/B//zPmTjttHPs00pflvDl8+15LJyHr331AvzwBz/F4IGjkJsrbYpm48w1mPTCaCwZI+Ivf"
        "ipKw6LirzQaVrC5LeexCZMmT8eVP/0NLr74F5jzwVJUhk9bBl+AAgoKijFyxAR8w+z+L//xLzHvoyXOI5nc6VfbXT/FX/71K5HGYxlHMeuVKeh0xVP2Q+BZb8xC24efwL8ffw7Llq4KcwVw8x2UdEOU+MuTR7UpLCrCvHmL0LVrHzz2"
        "r+fxz3bPBeHR59Du0WfRrl0Q/mXSHm37LO74/f340pfONq9/RfrqDaGVaNZOXYHx7Ydj6ei5dqGKtVBRGhYVf6XRCD5UDcQrY9c+tHmkPb7whW/g3HN+hGfbd0H3bn3R3Qhoty690aNbP3PeH61bP4XTv/ZtI8Cn4/HHX8SePfKsfSC"
        "D3OlXcucfLiokf/cxzO0xA12vfQHD/9YX2Uv3IG3+Gtx6Y0t83dj6y58fwqqVa8LcgUeBV/xr307Ys9riz9s+7s6/NhUVVfbzi927M+0iERXo+/btu/H22Kk44+vfxk9+cj0+mrc4tBCBWedWjFmAcf8eglVj5Z5/4JuiNDQq/kqjIU"
        "/U8A3Am6PG47LLfoXPf/6b+M63f4qrW9yMq35+PVrYcAN+8fMbcVWLm/Czn12H8y74ET77ma/h+9/7Oaa8OxNlZcHu30q1+VNpd/zhu4WsPMzvMwu9ft0RA27rhk3vpwOl1Th2pAB9eg3FRRf9BF/+8nm4797WSFu9Lvi8wBCTVB4FZ"
        "9n7+YEv7/nLB77uc/5BHr474e596ZI05GQfsnF1sXx5Os4592Jc/r/XYvbsBbV+NsKlvKAc8wfOxLh/Dsb6qcvDWD+XojQMKv5KoyHiv2XzDtzzpwdwzjmX4Nbb7kOHl3rg2Wc7mt1/BzzzTEc80/5ltH/aBHP+3PMd0f7ZTvjVNXfg"
        "i//xTTz80OO2fICxZ3+TJ7BbklOEJYPnotcNHdH7Ny8jffhSVBXwXUFAzoHD6P7aAHzr2/+LL/33Obj/7+2wOn2d/a2dKKz4/zF42qdt2yexu2bnz4UmWGy2b8tAu0efxr1/fhjvTf3AxvG2FX/nxwbzjsAGflhtHCktLcdHHy3FWWf"
        "/AD++8jrMnrOw5svL9FN8Jfn78jG90ziMf2wIds7fGMa6ORSl4VDxVxoVCmHvnoNw4QWX4te/vg1zP1pixLfaPh5ZUlxqXstMKA9DKUrLypCXX4Qxo6fi7LN+iPPPvQxDBr1p0wg/6EV1BUoOl2DVsMXoa0S/xzUvYEHf2ajIDX5/P9"
        "RWS86Bg+jWvT/OOftS/NcXz8LfzQKQnr6+5ikjFz7qeXfL+/GZz/5fPNLmiRrxDz67CKymm3cPV15xDc78xnfx9JOdzK5+DZYtWYUli5ebsAJLFq3E0kWrsHTxSixfloaFC1egX78ROP1rF+Lqa27DwkWyow9k3ZX2rLS9eOexoZj0z"
        "Ahkrw8XnnChU5SGRsVfaVS2bNmJ6669HeeefTE6vdyj5rdtEmL0jt+O/cufW+GLX/g67rzjPvvFMFKFclQUFiF94hL0uulldL7iKczpPBWFZtdMrFbaTXdMWrOyctCpUy989fRv4Yv/+Q089OBjWLt2o/MoaUBwz/9+fO7zXz1u589/"
        "WpHsN7b+el9rfP2M7+CCC36Eq6++GVdfdROubnGjCXy9GVe1uMWGq6+6GS1+cRMuvaQFvnTaWXjwocdqfcO3xkO6Ycyvm7wSIx7sh5k93kXBfvkyWNAGRWloVPyVRmXGjDn4wx33om2bp7Fyxdow1ggaVdrqGv9QWMOnd0Kt4/H0aR9"
        "Y4efv7UyePDVIr6pG7tZsvNthFAbd1xMf9pppzoN777zhY7/7W2WOTL7wISILHzPt1rUffnnNb3HD9X/A4EEj7Td6XXIPHcVLL3Uz71B+i759BiOnJj0m/nwTsHDBCjzx+Eu4/ro70MII/lVG4K+yoh8T/qta3GqEn6834brrfo+2j7"
        "bH4iVpKA5/TZSu1bhnXK44WoYZr07EyAf6I23cclSV1HRE8KooDYyKv9JoUKz378/GmvQN2LljN4qKSsIUHwrc8SKXl5eHTRu3ID19HfbtywoeCTX/8586zNm8D9nrs1CQXYCq8kBGg91+cGT/mhfRTu7yDxzItb6sWLEa27fvQpF5B"
        "0GCRSf4vf9du/bY+rhYRN0aIiUl5cjadwAb1m+xTxGtrBXWOmENVq1ag3XmXca+zP0o5bd87XOm/KGK8N8ONvYYDqZnYuSDr2PcEyOxZ0X8H39TlIZCxV9pVCjY8nRLfaAYC7Rh7Vg1N/9XGNE0IdT4eiMf0Lp1BASLQBDCqLqgG3WE"
        "Wtg4vj8J6mGGqtIqLB72Ifq17II5faYjb598+Y25FKVxUPFXGo3jxTV54pUNBDM8MdSI6ElCGw1hpy6sv/a/ELMwHt12ECPbDMKQf/TBpplr7a97homxfIrSwKj4K41GYwhqzCYDRTI4T1SPmxYvn9jw0+PlJ1H5jydm1wY5D1NLj5V"
        "g9VsL0Ot3nTGj22Qc3infHWCOWD5FaWhU/JVmhYhoIPyB+J9aBKIdN/AT4/A2WFVZJTLTMvDWw/0w+C+vYfuHm1BZHHwHQfb8wV9FaXhU/JVmSiimzZi8vccwp/c09LjtRcztMw2FWcHjqnZJs2ub7vyVxkPFX2l22J2/DWFEM4KaTs"
        "qOFWPduOUYeGdXjGg1EPvXZqKqnM//WM0Pxb9ZNlFpJqj4K6c83P/KHpi/h19eWg5U8ukfRxx5IOFjxFZv/eDtKT7OyYc6TUQo5qSqrArbPtyENx8ciIF3d8GayctQXhh7rDRY2/gnjFCURkDFX2lWlJvdcf6RAlRVhI9LWoU8xVTSu"
        "uSIP3VcXDQLQOaKDEx46i30/n1nzOk9BYVZx2LpitJEqPgrzQbqY1lZObIzslFSWG7Ftebbwfa/UwP6Yf0x/vE7DjXCbs6zN+7D+x0noN8dr2LCsyPNeaZtgi2jK4DShKj4K82CQNqrUVZejj3rM3Bgew4q7L+LS+Hnr/y77wRCgiJN"
        "TOCD9YZizm0/MQJ/aEu2fZyTwj+m7RDsXLDJxAceB8Lf5M4qKYyKv3LKE8hi8LeyohL7N2Vi7dRlOLonN/imrk2j3Aahho9FTwMfgn9xjJiz8koc2ZqLOV3eR9/fdcaoVv2xacYaVIVf5uJfFX+lqVHxV055XFnkbZTCnDzMG/weNkx"
        "PQz5//bImkf9bKa05r0lrKngrird77H9mY19aiZyNWZjeeRJ63/gy3nygPzYb4a8sCn5+2rpnhJ8lFKUpUfFXTnkCIQ3ElJSXVWLjjDR80H0i1r+/GkW5wQ+0BTBXuPv/GPQ05iVQUVCOjAXbMfGpt9D95hcx+pEh2DJrPUqPBj9wZz"
        "+z4EJh/z1iG6UoTYaKv9IMcMXfvJqD/IzDmNZ5PMY8MRQbZq4GwiclefuEt9nt3tuKqzlhwXqGiKgY5oTnwS9z8nmewKcazHH+3mNYNXYZRj04ED1v6oCJT76JjIXbUJYf/KQzCT6wlqAoTYuKv9IMcKU/OENZNTabXfTIRwbirX8Nx"
        "pbZa1FdGvu3fq0Y8yegZQEIS9YOdeHktYtI7JyLjA1W+oPf+iflR0qxY94mvNfxHfRv2RUD/tgds7pOxf7Ve1ER3uqpjdhUlKZFxV9pBsTEkUfca5PS3GIsGTYPg/7SE6MeGYQtH6xDubezDoIILMXa/U9ijw8WpxyPY2If2HXhLZ6c"
        "DfuwdNQCjGw9CH1+1xGjHuiDpcPnIXfHIftUj6KcSqj4K80AEeDgiM/3VCHY5eftOowPe01D3z90waiHBmPDu6tRkHPMfjBMgr8kkG2eJxOiToLyNZGWciP6h7cfxsYpazHl+bfR5+5X0feuruZ4LLbPWmcWqPAfjIlYMBTl40TFXzn"
        "FodAGX+Sy98bNqd17V1P8AxHO3ZaD2a+9jz6/fRVD7+mFxUM/RM6mbJQVlAXFQgLZDv+6u3o3uO8S+OLDLBVVxnYJDu86iE0z12N6p6kY3LI3XrvxRQy7vzeWvDEPh7eY3X54l8cuGTXvQBTl1EDFX2kGUMEp9rEFIPhUlweBoB7LPI"
        "xFQ2ZhwB2voIcRYf5ziBsmp+Hojlx7S6aaX6ai+FoB5vIR2gqWEhuCX+FhnJVr+5/935Rh+crCShTnFODAmkykj12KKU+PwcA7uqH7dS+i752vYGaXychavBNVR51bT/LH2uGJopwaqPgrzYBQPW0Ixd8GI8p8Jeag+FABNk9fg7cff"
        "QM9b34JvX/bEaPbDsHCgR8i46PtOLztEEpyi4IvV1HjE2HWBv62fnFuMQ5vP4TdRtTTRi/BlA5jMeTenuh540voZcLw+3pjdvd3sWP+JhQdLLTvCmIOcrfvLiaKcuqg4q80CSKHAcFZ7bgESEYTakuoCG2IOSzLL0N2eiaWDJ2L0W2G"
        "os9tXfDaTZ3R/4+vYey/hmFmt8lYMmIe1r23Cts/2oTdy3Zi78qMMOzCrsXbsXH2OqwevxSLBs/BzFcmYty/3sCQP7+GXrd2NILfEQPvfg0TnnwTS4bNNfm3oiDrWPCPsDgLSiD2jHAXA0U5dVDxV5oYiqDshkNJPGFdrLFgkYWhurw"
        "ahdn52LNkF5aPWIgpz47F0L/2Rp/bO6HnrR3Ma2cMaNkdb9zbG6P+/jreenCADaMe6I8R9/fFkHt6YcAfuqHf715BfxOGtOyGkQ/0xcRnRmFe31lYPzXd/v5+4YF8VIaPl5LoZpxw4xSlUVHxVz4GQpG2fw1J3wxnPgkxeMblJPjtfG"
        "f7bXSZt3kObNyHLTPX2HcD01+eiPGPj8KbDw/C8L8ZoTc7+sH3dMew+3pi5N/72Z9feLvVULzTbgSmPDcW83rPwOqxS7Ft7gZkr89EXlYeKgr9XT5Pg72+ojQXVPyVpsHTxeDU/LXCTyUN3g1EwdggJZB4UV6JZwgsBB/dBs/jByVq4"
        "D38vArk7zmGnA1ZyFi6HVvmbLC3f9LfXYb109Kw7cMN2DF/s3nHsBNZ6XuQu+Mgig8Vozr2+W1tbOVx6lOUUxwVf6WJEJkOZDKQb1fM5fh4EZWSwd8o8Td/uYjY/51nd+zjlbG8J0NQT/An+MVOnoX2w7oVpTmh4q+cFMk/u858RjJN"
        "fj6hGYhyTJhFsAN7zCfHtbG77DDYn24IbdRgovlPPdry3kLh5QwiPRgV5A3fRdizMNKeSUoQbIIUUpRmhIq/oihKCqLiryiKkoKo+CuKoqQgKv6KoigpyCdK/E877TR86lOfwnXXXRfGxMjNzY2b/vbbb9t4hmnTptk45uF5ly5d7Dl"
        "Zs2YN7rrrrho7DDxnfLIwP8vRBn06lWHb6WtUf54oCxcurOm7eEi9Evx+atWqVa10pWFgv/P6JKIh5oAQNceShfW1b98+PEtuXCm1+cT0lCvgDDt37gxTYlDYJV1E3l0UOIgFf2DSnjvg/VDXpCG04Zahz6cyp4r4+/10/vnn10pXTh"
        "4Z74nGcUPMAZcTFX/OXfrhjksV//rziekpGUgiDNwdRsHdAtM5eCj88Xbi/sAUQaJ9ycfXyy+/3Ma7C0c8XBvymmrUR/yjriV3fFJegnLySF8mEvCGmAMuJyr+4ocr/kr9+UTMHHdHLQPDF3OBcSIqMmgZ5J2AkEj83QnCunnO17qQe"
        "sUWgz/ZpF4uUu7tDR677Uk2H+OYl3XKrs1tk/jEwP5w+0H8dCfZwIEDa5Xx3/LLbQFJZ52uT/URf3ehFiTNvXZCfep228FX//rX1TeEeaRPxYbkd3FtiU/ueJFryXjXf7afbZK20ob/Loh1un3B8u6Yimdb+sXtEwnu9XaRvqcfbh1R"
        "c4C2WYf0j38tiPhGu0Tsu/W7/hEp4wbm8fMJbt8zsA/c8Sp1sg/Z3+Iv63Hz0b7bz1HjobkRfwY2I0QkeJE5uOQCxbut4g4UBg4IH39gcmDLwGDgMctRRNwBHQ9XGGhLBhInhIs7uFkHz6VelhGSzSd53MB+YeCxlHUHtkxifzK6Zdj"
        "n4gPPpQ9kotGe65O0M94kdZF6+Sp+sRyRc8nj2qlP3QzM4wqDtCGZvpExx8B01w6DwLp57tty+4zxUo7pbn1Rcb6fDLTh2pH+cuPoo3sui4tfP+OjSHYO8Fj8lXZLOcb77ZY55o834o8X+iZ9LbbZhqhxJX3PwHxuORF2dxxF+UrYbs"
        "lDm65d6efmSPwZ2IyQi8VBSOTi8GJHwcHnTtaoC+gPTMIB4+6eJLB+TsRESDlZaNyJK4JCpF76J5OE9Upe2W0km0/OZdJJXqazbTIJiOSV/vAno5zTlpRjXvGftplHrgPxbURNUh8pI7Z4zElP+zxmf/t26lu3XC+xySDtTqZv/DFH3"
        "LFBRDSYV/qdyLWTsnIuYkN4zDjxPcpPGcNyTqTNMs7EtjsXZOES24Tnvq0o2Cd1zQG5Zoxzx4bfZ+IbfSb+9SL0h3EMQjL5ouYCkTqlf8QWg/jK/BJHxDb9F1u0746P5kisR5spUSLqXryowSyDX4I76QR/YLpwINOuvwuJh4gAQ5To"
        "uHVIvbTt4vuTbD4eM0i9Lhy8zMeJIGLDIH3GNJ7TJnHbIYHl3MnFdrEu+sVyMuHFhj9Jo5B6+SqTmP0s15r+Rtk50bolTtpNEvWNa4d1Cr59EcF4Qa4d/eM56xOi4qQc65F+iRdkTEfZ4THjpF+IlHP7IBGJ5oAsDv7Y5DnjRXh936L"
        "88vuUJJNP+t6f26IN4muydUobJfBc/G6u1J4FzRAZQPEC3wW4uAuDXHgGf6D6A5N2GOeLqDtQeByFW09U4EASpN66/Ek2H48ZfN/EJ04ClpGJ6eaVPEwXKDrsC38ysAwFQQSXk46TXPwRG25/xUPqlTZIXfLKa+DbOZm6Jc5vd7y+cY"
        "U3kfiLHfrNND/QDhE/pb0kKk5sS3k555h27UogUXbEL+kXIrakXBTJzoFPovjzOtN/2pM0BvG9OVJ7FjQzonaiUUHeEbgCIRdbLj4DL7rgD0wZ0JzIMmmJW17q8fGFMirIhJJ6WUaExRUb2WUnm0/O3bYR6QfJR/y8/sSgj4yTMqxXJ"
        "gPjoyacbyNqYvlIGb4SuY0ngfX6dk6mbomTdifTN5JHfCQyRhiIez3cMcP2sJzYkmvp2oqKE1u+D2y7wGOKlD+eXDs8Zpz0C/FtR5HsHJBrQf9kbDLe99f3TezQvuDaFqL8p99uPrfvpU30ReqUBSgZWzxnPrefZUy65ZobtWdBM0Mu"
        "AAeVj7sw8MIRGbzM7wq1CJgrpP7A5GCSwRsV/F2OwEkoedwJI7BOpskgknoZWB/PpV5X2JLNJ3n8SS1t5ivbKH4wiOj5E0POWQ/73hVl9ifLyTn7w00Xn/yJFYXUI33v2o1n52Tqljjpo/r0DQP7R8pIEOQ6yTUSWzyPN9ZIVBzPGcT"
        "PRD5Inig7Uo5pgvjFVxFGn2TnANvl+iI+MDA+0Rxz88XrU3dO0QavSdS19ceA+M5XmYtRfeHbcs/ZN2yn2JJFtjlSexY0IziA5AK4A9tFLj4HtCsO7gpO3EEnA9gfmIQCR5vuBGC+RANA7LiDy8UdyLQv+emHLFYMPJZJQ5LNJ/EiBg"
        "LbLDYY2C7pL76SqInBOFcMOalc20yX/mE+v31Rk9RH6nX7XspIXJSdE61b4qQdyfQNYX3SF+wH2fEyCLwWbj76Rxv0R5C63PZGxYltt7/ZRlckef3d9Cg7PGace11px+27eCQ7B9huN5+0O2oM+75JGbaL41vqEWhDyjKw7njXlvak7"
        "xnYPyL8JKovomxRP9x+pk1fR5obtXtK+diJmhBRJJtPaRwoOBQEV8QZx2tCkVCUUx0V/1MMFf/mgexIuUvltZDroddEaS6o+J9iqPg3H9j37i0F7vj1eijNBRV/RVGUFETFX1EUJQVR8VcURUlBVPwVRVFSEBV/RVGUFETF/wSQR/tc"
        "+MUT98tW/hda3C+OuMH9Qg4fH5QvuNC++2UUHz5V4tuS4Nr0oU/+Uyr+F1+Sgc+0u2V47H/RpzFg29x+Zjsa+gkb6aOGRp7QcgOfEJJvDRPW64+t+sCyJ+o7213XF7PqwvffHycnQ1Tf+P3pBiUx2kP1RL7F6Q9CTmKKEr/0w8HOc04"
        "cQcSfr26QiUXhp5CxrHyLkufx4ERgHb4916YP41mGdjkpmZfCw7ZwotdnkkpbhJMRnWSRL1Gxb6StrFNEqqGIEpmGgDZd39n3spBJX55s3SdzHfyxwVeeMz7ZBcD33x8nJ4pce79vpC8lsE85HjiflMSo+NcDiiMHFoM7CBnPgel+25"
        "ODkHEyaeqa1LTJAS6wHMu7u0KXExEJCg/riZrInOD1sedPapY9UdFJBvYt64yqQ/ra9edkOJG+TYZ4fUSBlfiTrftEr4OMYX+8sU8Z747N+nCy14VjlQskxy37qa6+YV6OZaVuVPzrAQeVTE53EPLdAAenjzvwuROJN3Bl4vk770QTu"
        "b4iIYtJvN8jYd3uJKUIsL0sw7bRf9oQMZBAH9xzBoE+csIyjq9u3eK/f6sramEiki8e9FcWX9qQhY522Q63bYyjL277RNz89ki5eP1B3LYwnfmiYJ6o6xlP/FmPf71Yh9hn/fLOQXyKV0ddyOIaNT7YB9LWKPtunPgfNU5IXdfGh2ks"
        "S//4KnaikDoT2VNiqPgnCQe1TDp/EMqA9+FAdEWFwRVD2WXJoPVh/ni3M8QflnWDv4AIUgdf60LyUgg4WXnOuig0gm+LvooAEB5zgksb2Q+umDGdNuQ2CPOxT9w6XKT/koH56C/7gv6LL7I4sF7xjXWLaEvfMb9bV7z+YDkibWEZsRk"
        "F06W9EmRXK765dbtCL7iLKNMYWJaBtukHbZwIUp42aSOqHfTNt+/Guf4T2nPtMI32412bRLCsa9uHafHmi3I8Kv5JwIHKQSzi4A9Cf8AL7kTkpOWgF3GQiSYiwGMfvx4X2mUZP8TLL3W4EzEeFCR/EtFvlncF1LXFeqWthBPav1VAMW"
        "M/EPHfhXEsF0WivnCRa+WLCftexJrprq+Ece7C5NaVqD9EwOQ4EbTJfH5gvJR16/bb4p7L9XTbSRvsP79t9YHXTBYk2vft0Tffvhvn+k9oQ8ZJMtcmEbTr2naR/pA5qtSNin8dcEK5uy3iD0J/wAscjL4AutAuRUUGrg9txtvJxKszH"
        "lKHTMREMJ8/wYlb3rdFX6RMvLoknkT5zzhJ92HeZNrL60QbUUHK89j3jXHiv++bCGFUoJ2otkTBPFKHwEWE4sdAfFscI1KGY0HeGTGO9ftE1XGiUEgpyqxHbEbZd+P46vovfUSSuTaJYJ54+dg30odKcqj414EIVrzAdA7qqB2rpMdD"
        "BrPsiPxdC9PiTWR/ktUFFzHW4S5iLvSTwiL5oup12+O3zfWV8X46kXgS5T/jJN2HIhTVxwInPxfaRDaEKN8YJ/77vrlpUfj54+H2kYv7rsq3xesl75bYftlMiJD6xKujLtgfsrD4sG/Fpyj7bpzvv9vXTIvyOVlo17Xt4r9DUermxK9"
        "EikAx5OB1g+zUeMz0qLez7m0BpnNw+m93Oanl7a47sYmIMO1E4U+yZOAkZj207cP2iMgkus0hbeAx2y+4AkD89hC2VeqI8j+ROLBepkVNcPGN9cm1cH0jbI/0ZVS6a9v3jf3h+8ry7i48mWvh95FAv1k/r4tvy203+1SIGnPEfadQH6"
        "QP/X4h7DtpK33zb9G4dfr+uzaTuTaJoF3XthDPrpIYFf8TIGoQ8pwThJORgWLqCqgsGJLONHdB4KThJOJApggwXYQyCuaXBcgPYtOHdlkn7VJwmJeTjr4znnUTxnMycXfJMjxnXSIAhOn0QcpI+5mXMI02ZVKzPp7LOw9fJAjjaDcek"
        "s6+YT0MUo/b17QrfUP/5V2D9EuUUEh7CF/ZR1Ker5Iu50wXEWS835YomMf1nYH9Qd+kfJQt9qvfRiLt5DWQMeW2g+k8F3js23Zhmlwj8U9s8phIX/K6uulSp++/pLnjJNG1SQTLRvkvi6dSP7THToCoQciBzEnKQcjAScE4gYPbTWd5"
        "mRCCTIR46S6cUGLLDzIRo6BP9I3ixbysj375dXFyc5JKHvrmtkfqlwWK+cV3mcjMI/XwVYSfMM3vQ7GZCFmsmI+BPvrtlTbG60vGiZgJjBM79F/8prAQvz/cOqPaEoXrtwTacsdKlC0RN99naafY4rHbH1KfwONEftIey0rbGdwFnTC"
        "PP87dOn3/ec58Mk7qujaJYN4o/8VnpX6o+CtKiiA7dUUhKv6KkiJEvcNTUhcVf0VRlBRExV9RFCUFUfFXFEVJQVT8FUVRUhAVf6UGPg0ij/BJkEcR3cc3Ez0u2BTQL/+xx5NFHuP0nzeP6hM3xPODH6yynySf/2FrPLuuvWQf/eXjk/"
        "7jp4SPiMpjm3zl46qKIqj4KzWIIPFVAp/Np3DIc9SfRPGnqNImBdoXUAqr2x8S2AcUXH+xICLGsmgy8Nh9Fp02/L5mYFlC4Wd++hZV3kWeu3d9p9Azjq+0yevI8yh/ldRExV+pQQTJR4SE6aeC+Dc0FFoKqOyU60L6Q74A5iP9KEJOJ"
        "E7Et65+5OLh2qctlvd378wj7w5c8Wd7/Gf6+e6GbVUUouKv1BBP/CXeFX++iujw3BU67lZlN8rAY3fHybwUJilPUaJtgXHcqcqtGF8IGcf8tMk0V/TkVonrT12IfZbhsS+wLlKnL6wuYkf8kvayPQL9jCf+8k7Ev80j/S6IL3J93DTG"
        "u9+oJkxPtOAoqYWKv1KDiIgLBYaCQTGhiFFAmIdixvwUSu6WKfBEhI/pLMvANOZhGqE9CiHFTWyyDPMS2uc5bbMOCqUrhjxmPJHbGUwT0Uwk3j7My7oE+p1I2Ok3g7QlHvSFdukPg9t+wj5gYLyki99R14Ewv+sbz2Unz/yu+PNc+kh"
        "gOutRFKLir9QgouMHCoYIr4i/i4g3kc8IXCh6LMPdtQi0CL1AQY0nZIRxspPlsStsIqS0IYtQsjC/1EvEvyhx9xeheLCs+EI/GWTRENhHPJcFjqIu7WLgsY+0k0g/i5/Mn4z4R9lVUhMdCUoNIjoiQAy+0FFARIAEV1RcgXJhHPPJTj"
        "0qSDke+8LFOBE3P50LCRcfhijRjocsSlFBFhqBIh0VHwXz+L7QR5ZP9K6EYs5FgG1jXh/2D9NlgXKvDc+lfwjP/T5kur8wK6mLir9SQzzRcaGA+OLOOCnHND+dMI753LzxiBIuxrGsHLvp4jdDXbtyF4o0d98s7wbu8F2RlMUl2XcVU"
        "X1E3DZEIX0XJe7E78N4gdBff6GK55eSmqj4KzWIiCYiSkBEjAgFx99dyg7bve3Duly4o5VdcVQ641iPHEs6bbM+plGc3VsrdcG87i0fQXbpUgfzsY5k31Uk0wcUZ//WF8uIP0x3P+SW8uwjlqNvbmCavGsg7Aueu8Rrr5KaqPgrNYiI"
        "JKIu8adIUbgoPBQpBgqRK54sLztuxlGQXDGkLRExgXFR4i+7dNphedqRfOKXb4vE210L9I9toH3mo6DTjh/EZxfxw+0D2mMQ5FzSmdftA/rOdtE/to3p/oLi4vYPoW+MkwWV/vNc7CuKir9SgwhGIigwicSfULAo+Ixj4LErOiJmFDu"
        "m054rwoyjLy6ME3GTdPHXzSsiR3vil2+LyKIRD7FDgeZrvOAKrovfB2yvLH6E/eGm+31AZFGMl+7CPL4vfOfANjKNr4k+b1BSDxV/5RMNRVN3u4pyPCr+yicW7vi541YU5XhU/BVFUVIQFX9FUZQURMVfURQlBVHxVxRFSUFU/BVFUV"
        "IQFX9FUZQUJCnxv+aaa/DpT39agwYNGjSc4uGGG27AxIkTQ/WOj+78FUVRUhAVf0VRlBRExV9RFCXlAP4/um3BQcY5cnkAAAAASUVORK5CYII=")
        #image = "./AXIS_Logo.png"                                          # Instead of using the base64 string, you could get rid of the 'axis_mailhead_base64_string' var above and the line below, then add file to pyinstaller dependencies & open file / insert into word doc
        #image = io.StringIO(base64.b64decode(axis_mailhead_base64_string)) # Convert the base64 string into file type object to insert into word doc (this works for string conversions, would use this for python 2)
        image = io.BytesIO(base64.b64decode(axis_mailhead_base64_string))   # Convert the base64 string into file type object to insert into word doc (this works for bin conversions - b64decode returns bin code in python 3)

        # Pop the spreadsheet's title row before we enter loop
        borrowerAddressArray.pop(0)                    # Remove the first element from the list (or the subarray that makes up the first row in outer array)

        # Create rows and add data
        while len(borrowerAddressArray) != 0:         # Loop until all mailing addresses have been added to word doc (addresses are popped off array as they are added to word doc)

            # Variable Declarations
            counter = 0                                   # Create a counter variable 
            table = document.add_table(rows=0,cols=3)     # Create a table with zero rows, but will have three columns when rows are added

            # Create rows and add data
            while counter != 3:                           # Loop until a table of three rows has been created

                # Add a new row to the table
                row_cells = table.add_row().cells         

                # Modify Paragraph contained within the first cell/column of the row. 
                if len(borrowerAddressArray) != 0:                 # Only add the borrower details in the table if there is one left in the array
                    paragraph = row_cells[0].paragraphs[0]         # Access the first paragraph in the cell at position 0 (think of each row as a zero-index array that is accesible by elements from left to right)
                    run = paragraph.add_run()                      # Add an empty run to the paragraph
                    run.add_picture(image)                        # Insert a picture into that empty run. Note: The width, height vars of the picture are optional. Ex: run.add_picture(image, width = 1400000, height = 1400000) 
                    paragraph.add_run("\n")                        # Add a newline character
                    paragraph.add_run(borrowerAddressArray[0][0])  # Add the borrower's name
                    paragraph.add_run("\n")                        # Add a newline character
                    paragraph.add_run(borrowerAddressArray[0][1])  # Add the first portion of the borrower's mailing address (subject street)
                    paragraph.add_run("\n")                        # Add a newline character
                    paragraph.add_run(borrowerAddressArray[0][2])  # Add the second portion of the borrower's mailing address (city, state, and zip)
                    paragraph.add_run("\n")                        # Add a newline character
                    paragraph.alignment = 1                        # 0 for left, 1 center, 2 for right
                    borrowerAddressArray.pop(0)                    # Remove the first element from the list (or the subarray that makes up the first row in outer array)

                # Modify the text contained within the second cell/column of the row. 
                row_cells[1].text = " "

                # Modify Paragraph contained within the third cell/column of the row. 
                if len(borrowerAddressArray) != 0:                 # Only add the borrower details in the table if there is one left in the array
                    paragraph = row_cells[2].paragraphs[0]         # Access the first paragraph in the cell at position 0 (think of each row as a zero-index array that is accesible by elements from left to right)
                    run = paragraph.add_run()                      # Add an empty run to the paragraph
                    run.add_picture(image)                        # Insert a picture into that empty run. Note: The width, height vars of the picture are optional. Ex: run.add_picture(image, width = 1400000, height = 1400000)  
                    paragraph.add_run("\n")                        # Add a newline character
                    paragraph.add_run(borrowerAddressArray[0][0])  # Add the borrower's name
                    paragraph.add_run("\n")                        # Add a newline character
                    paragraph.add_run(borrowerAddressArray[0][1])  # Add the first portion of the borrower's mailing address (subject street)
                    paragraph.add_run("\n")                        # Add a newline character
                    paragraph.add_run(borrowerAddressArray[0][2])  # Add the second portion of the borrower's mailing address (city, state, and zip)
                    paragraph.add_run("\n")                        # Add a newline character
                    paragraph.alignment = 1                        # 0 for left, 1 center, 2 for right
                    borrowerAddressArray.pop(0)                    # Remove the first element from the list (or the subarray that makes up the first row in outer array)

                # Iterate the counter
                counter += 1                                       # Increment by one

            if len(borrowerAddressArray) != 0:                     # Only add a page break if there are more entries in the array that need to be generated 
                document.add_page_break()

        # Format all pages that have been created 
        sections = document.sections               # Access "sections" (aka pages)
        for section in sections:                   # Iterate through all sections/pages in document
            section.top_margin = Inches(0.50)      # Format top margin   
            section.right_margin = Inches(0.25)    # Format right margin 
            section.left_margin = Inches(0.25)     # Format left margin 
            section.bottom_margin = Inches(0.25)   # Format bottom margin

        # Save the document & overwrite the file user created with the changes above
        document.save(file_name)
            
# --- main ---

if __name__ == '__main__':
    root = tk.Tk()
    root.title("AXIS' Mail Label Application")
    top = MyWindow(root)
    root.mainloop()